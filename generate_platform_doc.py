"""
鹰眼工业智能运维平台 — 完整参考手册生成器
Comprehensive platform reference manual as .docx
Every chart, card, table described with: what, why, how to read, how it's generated
"""
import os
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

OUTPUT_PATH = os.path.join(os.path.dirname(__file__),
    "鹰眼工业智能运维平台_完整参考手册.docx")

doc = Document()

# ── Page setup ──
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.35

for i in range(1, 5):
    h_style = doc.styles[f'Heading {i}']
    h_font = h_style.font
    h_font.name = '微软雅黑'
    h_style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    h_font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)
    sizes = {1: (22, 36, 18), 2: (15, 24, 10), 3: (12, 16, 6), 4: (11, 12, 4)}
    sz, sb, sa = sizes[i]
    h_font.size = Pt(sz)
    h_font.bold = True
    h_style.paragraph_format.space_before = Pt(sb)
    h_style.paragraph_format.space_after = Pt(sa)

# ── Helpers ──
def P(text, bold=False, size=10, color=None, align=None, sa=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '微软雅黑'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if color: run.font.color.rgb = RGBColor(*color)
    if align is not None: p.alignment = align
    if sa is not None: p.paragraph_format.space_after = Pt(sa)
    return p

def RP(segments):
    p = doc.add_paragraph()
    for seg in segments:
        text, bold = seg[0], seg[1] if len(seg) > 1 else False
        italic = seg[2] if len(seg) > 2 else False
        sz = seg[3] if len(seg) > 3 else 10
        col = seg[4] if len(seg) > 4 else None
        run = p.add_run(text)
        run.bold = bold; run.italic = italic
        run.font.size = Pt(sz); run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        if col: run.font.color.rgb = RGBColor(*col)
    return p

def B(text, level=0):
    p = doc.add_paragraph(text, style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0 + level * 0.7)
    for r in p.runs:
        r.font.name = '微软雅黑'; r._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        r.font.size = Pt(9.5)
    return p

def T(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, h in enumerate(headers):
        cell = table.rows[0].cells[j]; cell.text = ''
        p = cell.paragraphs[0]; run = p.add_run(h)
        run.bold = True; run.font.size = Pt(8.5); run.font.name = '微软雅黑'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1A1A2E" w:val="clear"/>')
        cell._element.get_or_add_tcPr().append(shading)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            cell = table.rows[i + 1].cells[j]; cell.text = ''
            p = cell.paragraphs[0]; run = p.add_run(str(val))
            run.font.size = Pt(8.5); run.font.name = '微软雅黑'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)
    doc.add_paragraph()
    return table

def CHART_BLOCK(chart_id, chart_name, chart_type, axes="", series="", data_source="", purpose="", interpretation="", generation="", *_extra):
    """Standardized chart documentation block. All params after chart_name have defaults for flexibility."""
    if _extra:
        generation = generation + " " + " ".join(str(e) for e in _extra)
    P(f"▌{chart_name}", bold=True, size=11, color=(0x00, 0x6B, 0x5E))
    if chart_type: RP([("图表类型：", True), (chart_type, False)])
    if axes: RP([("坐标轴：", True), (axes, False)])
    if series: RP([("数据系列：", True), (series, False)])
    if data_source: RP([("数据来源：", True), (data_source, False)])
    if purpose: RP([("图表用途：", True), (purpose, False)])
    if interpretation: RP([("如何解读：", True), (interpretation, False)])
    if generation: RP([("生成方式：", True), (generation, False)])
    doc.add_paragraph()

def STAT_CARD(card_id, card_name, value_desc, calc_method, data_source):
    """Standardized stat card documentation block"""
    P(f"▌{card_name}", bold=True, size=11, color=(0x7B, 0x61, 0xFF))
    RP([("显示数值：", True), (value_desc, False)])
    RP([("计算方式：", True), (calc_method, False)])
    RP([("数据来源：", True), (data_source, False)])
    doc.add_paragraph()

def PAGE_BREAK():
    doc.add_page_break()

# ═══════════════════════════════════════════ TITLE PAGE ═══════════════════════════════════════════
for _ in range(5): doc.add_paragraph()
P("鹰眼工业智能运维平台", bold=True, size=30, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1A, 0x1A, 0x2E))
P("Hawk-Eye Industrial Predictive Maintenance Platform", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x77, 0x77, 0x77))
doc.add_paragraph()
P("完整参考手册", bold=True, size=22, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x00, 0x6B, 0x5E))
P("—— 评委查阅指南 ——", size=11, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
doc.add_paragraph(); doc.add_paragraph()
P(f"版本 2.3 | {datetime.date.today().strftime('%Y年%m月%d日')}", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
P("苗圃杯 · 半决赛作品", size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x99, 0x99, 0x99))
doc.add_paragraph(); doc.add_paragraph()
P("使用说明", bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER, color=(0x1A, 0x1A, 0x2E))
P("本手册按平台网页→Tab→图表/卡片的层级组织。评委在平台中看到任何图表或卡片，"
  "可通过目录定位到对应条目，立即了解：该图表是什么、展示什么数据、数据从何而来、"
  "如何解读图表含义。每条图表说明包括图表类型、坐标轴定义、数据系列含义、"
  "数据来源文件、图表用途和解读指引。", size=10, align=WD_ALIGN_PARAGRAPH.LEFT)
PAGE_BREAK()

# ═══════════════════════════════════════════ TABLE OF CONTENTS ═════════════════════════════════════
doc.add_heading('目录', level=1)
toc = [
    ("一、平台总体架构", "系统拓扑、数据流、核心算法公式、四级降级、三角色体系"),
    ("二、角色门（role-gate.html）", "入口页面 · 三角色卡片 · 选择流程"),
    ("三、首页 — 设备健康总览（home.html）", "10×10矩阵 · 运维KPI · 管理视图 · SHAP面板 · 故障注入 · Demo模式"),
    ("四、仪表盘 — 多维度数据分析（index.html）",
     "8个Tab逐一详解：Tab0业务总览 · Tab1运行日志 · Tab2数据探索 · Tab3基线划定 · "
     "Tab4预测性维护模型 · Tab5方案有效性验证 · Tab6智能维护决策中心 · Tab7设备健康与感知升级"),
    ("五、AI Copilot — 智能对话助手（chat.html）", "SSE流式对话 · 27工具调用 · RAG引用 · 图表/报告内联"),
    ("六、技术架构总览（technical-overview.html）", "系统架构图 · 算法公式卡片 · KDE数据天花板 · 流水线 · 演化时间线"),
    ("七、工单全流程跟踪（work-order-tracking.html）", "6列Kanban · 6状态机 · 技师分配 · 升级机制"),
    ("八、工作流调度（workflows.html）", "定时任务 · 配置管理 · 执行历史"),
    ("九、备件库存管理（inventory.html）", "(s,S)策略 · Gap图 · 风险环形图 · 采购建议 · 库存日志"),
    ("十、技师管理（technicians.html）", "名册 · 负载 · 技能匹配 · 增删改"),
    ("十一、知识库管理（knowledge-base.html）", "ChromaDB · 三级降级 · RAG对比 · 文档管理"),
    ("十二、报告中心（reports.html）", "5种报告 · HTML预览 · PDF生成 · 邮件投递"),
    ("十三、独立设备健康矩阵（device-grid.html）", "全屏10×10网格 · 详情面板 · SHAP探索"),
    ("十四、鹰眼球体3D数字孪生（sphere-demo.html）", "Three.js · Golden Spiral · 手势识别"),
    ("十五、评委讲解助手（assistant.js）", "话术库 · SSE AI回退 · 3D球体入口"),
    ("附录", "A:数据文件清单 B:API端点 C:技术栈 D:术语表 E:颜色编码参考"),
]
for num, desc in toc:
    RP([(f"{num}  ", True, False, 10.5, (0x1A, 0x1A, 0x2E)),
        (desc, False, False, 9.5, (0x55, 0x55, 0x55))])
    doc.add_paragraph()

PAGE_BREAK()

# ═══════════════════════════════════════════ 一、平台总体架构 ═════════════════════════════════════
doc.add_heading('一、平台总体架构', level=1)

doc.add_heading('1.1 系统拓扑', level=2)
P("平台由三大子系统构成：离线分析流水线（Python DAG）、在线Web服务（FastAPI + Vanilla JS前端）"
  "和AI智能层（DeepSeek大模型 + ChromaDB向量库）。以下为核心架构层次：")
T(["层次", "组件", "技术栈", "功能", "输出"],
  [["数据层", "原始脱敏数据集\n(agent-mcp架构/outputs_test/)", "CSV(UTF-8)\n4文件/2999行", "LOG/SUMMARY/ASSEMBLY/TESTS\n四类原始数据的读取与校验", "清洗后DataFrame"],
   ["流水线层", "五层DAG分析流水线\n(agent_orchestrator.py)", "Python 3.10+\nThreadPool并行", "数据准备→统计推断→ML→\n诊断→决策，全链路分析", "60+CSV/JSON"],
   ["同步层", "Dashboard Data Sync", "文件复制+CSV写入", "流水线输出同步至\nweb-dashboard/data/", "前端可消费数据"],
   ["服务层", "FastAPI Web服务\n(app.py, 8765端口)", "FastAPI+SSE+\nAPScheduler", "REST API + SSE流式 +\n定时任务调度", "40+API端点"],
   ["AI层", "DeepSeek + ChromaDB\n+ BGE嵌入", "LLM+向量检索+\n中文嵌入", "AI对话+RAG检索+\n评委讲解", "SSE流式响应"],
   ["前端层", "13个HTML页面\n+10个共享模块", "Vanilla JS+ECharts\n+Three.js+PapaParse", "数据可视化+交互操作\n+3D数字孪生", "用户界面"]],
  col_widths=[2.0, 4.0, 3.5, 4.5, 3.5])

doc.add_heading('1.2 五层DAG流水线', level=2)
P("流水线按有向无环图（DAG）拓扑顺序执行五个技能包，每个技能包消费上游输出、产出下游输入。"
  "以下详述各层的输入、处理逻辑和输出。")
T(["层级", "技能包", "输入数据", "核心处理逻辑", "关键输出"],
  [["Phase A-1\n数据准备", "data_prep", "原始LOG.csv(2999行)\n原始SUMMARY.csv\n原始ASSEMBLY.csv\n原始TESTS.csv",
    "1. 读取四文件UTF-8\n2. 按Equipment.Id分组(100台)\n3. 计算滚动窗口统计量(均值/标准差/CV)\n4. 构建成本风险矩阵(故障概率×单位成本×日产量)\n5. 计算逐设备Z-Score(V/A/T三维)",
    "z_scores.csv\ncost_risk_matrix.csv\nbaseline_stats.csv"],
   ["Phase A-2\n统计推断", "stat_inference", "z_scores.csv\ncost_risk_matrix.csv\nbaseline_stats.csv",
    "1. 复合Z-Score: Z_comp=√(wᵥZᵥ²+wₐZₐ²+wₜZₜ²)/√(Σw²)\n2. Hotelling T²多元异常检测\n3. 逐设备健康评分H=100·exp(-k·R)\n4. 告警分级(NORMAL/WATCH/WARNING/ALARM)\n5. 基线来源标注(自基线/混合/集群回退)",
    "equipment_health_score.csv\nalert_summary.csv\nvariance_decomposition.csv"],
   ["Phase A-3\n机器学习", "ml_inference", "z_scores.csv\nequipment_health_score.csv\n原始特征工程",
    "1. XGBoost训练(2999行×4参数,5折CV)\n2. SHAP TreeExplainer全局+局部归因\n3. 多算法对比(XGBoost/RF/IsolationForest/LOF/MLP/MTNN)\n4. 鲁棒性测试(噪声注入+特征缺失)\n5. 变体对比(窗口大小/架构选择)",
    "shap_dashboard.json\nshap_scatter_data.json\nalgorithm_comparison.csv\nfeature_imp.csv\nrobustness.csv"],
   ["Phase B\n诊断决策", "diagnosis", "前三阶段全部输出", "1. 多信号融合加权:\n   R=0.35·Z_comp+0.20·T²_norm+0.25·Cost_norm+0.20·Trend\n2. 根因判定规则引擎(电压漂移/热积聚/功率异常/复合退化)\n3. 技师分配(技能匹配+负载均衡)\n4. SHAP局部可解释性生成",
    "diagnosis_results.csv\ntechnician_assignments.csv"],
   ["Phase C\n决策生成", "decision", "diagnosis输出\n+策略参数(3选1)", "1. 策略选择:\n   成本效率(最少工单)\n   生产效率(最少停机)\n   质量优先(最低风险)\n2. 工业维护方案生成\n3. 备件(s,S)策略计算\n4. 停机窗口优化调度\n5. 技师排班",
    "maintenance_work_orders.csv\nspare_parts_plan.csv\ndowntime_schedule.csv\nprocurement_orders.csv\ntechnician_schedule.csv"]],
  col_widths=[1.8, 2.0, 3.2, 5.8, 3.7])

doc.add_heading('1.3 核心算法公式', level=2)
P("以下6个核心算法公式是平台所有图表和决策的数学基础。每个公式在technical-overview.html中有对应的交互式卡片。")
T(["公式", "数学表达", "物理/工程含义", "典型数值"],
  [["复合Z-Score", "Z_comp = √(Σ wᵢ²·Zᵢ²) / √(Σ wᵢ²)\ni∈{V,A,T}, w为传感器可靠性权重", "将多个传感器的标准化偏差融合为单一异常度量。权重大的传感器对最终得分影响更大。", "阈值: 2.0(WARNING)\n2.5(ALARM)\n精度83.9%, 召回39.4%"],
   ["Hotelling T²", "T² = n·(x̄-μ)'·S⁻¹·(x̄-μ)\nS为协方差矩阵", "马氏距离的多元推广。不仅考虑各参数的偏差量，还考虑参数间的相关性——两个参数同向偏离可能比单独偏离更危险。", "控制限: 99%置信\n精度100%, 召回22%\n(极保守)"],
   ["多信号风险融合", "R = α·Z_comp + β·T²_norm\n   + γ·Cost_norm + δ·Trend\nα+β+γ+δ=1", "四路线性加权融合：统计异常+多元距离+经济损失+趋势方向。权重随策略变化：成本效率提高γ权重，质量优先提高α权重。", "默认权重:\nα=0.35, β=0.20\nγ=0.25, δ=0.20\n(策略可调)"],
   ["健康评分", "H = 100 × exp(-k·R)\nk为衰减系数", "从综合风险评分R到0-100健康评分的指数衰减映射。低风险→高健康分(接近100)；高风险→低健康分(趋近0)。指数衰减确保高风险设备的健康分快速下降。", "≥85: 健康(绿)\n70-84: 亚健康(浅绿)\n55-69: 退化(琥珀)\n40-54: 橙色\n<40: 关键(红)"],
   ["Youden's J上限", "J = Sensitivity + Specificity - 1\n= TPR - FPR\n在最佳阈值处取值", "综合分类能力指标。J=0表示完全无法区分(随机猜测)，J=1表示完美区分。单传感器参数J≤0.075说明4参数条件下纯ML方法的物理上限极低。", "电压: J=0.077\n电流: J=0.064\n温度: J=0.065\n转速: J=0.061\n融合后: J≥0.90"],
   ["成本-风险矩阵", "E[Cost] = P(failure)×\n  Cost(downtime+repair)\n× DailyOutput", "期望损失 = 故障发生概率 × 单位时间损失(停机+维修) × 日产值。将统计概率映射为经济风险，支持管理层基于金额做决策。", "最大风险: CNC_036\n$7,232/天\n风险5级分档"]],
  col_widths=[2.5, 4.0, 5.5, 3.5])

doc.add_heading('1.4 四级降级保障', level=2)
P("降级机制是工业级可靠性的核心设计。当外部服务或模型不可用时，系统自动降级到下一层以确保始终能产出可执行方案。")
T(["层级", "名称", "触发条件", "可用能力", "降级影响"],
  [["L0", "FULL（完整模式）", "所有服务和模型正常", "完整5层流水线+AI对话+RAG检索+SHAP可解释性", "无影响，全功能"],
   ["L1", "STAT_ONLY（仅统计）", "DeepSeek API不可用\n或ML模型加载失败", "Z-Score+T²+成本风险+规则诊断\nAI助手降至TF-IDF关键词匹配", "损失ML异常概率信号\n融合权重自动重分配"],
   ["L2", "RULE_ONLY（仅规则）", "统计模型+ML+API均不可用", "纯规则引擎诊断(阈值+简单逻辑)\n工单生成基于固定模板\nAI助手不可用", "损失统计精度和ML信号\nSHAP不可用\n诊断粒度变粗"],
   ["L3", "EMERGENCY（应急模式）", "数据文件损坏或全部服务不可用", "硬编码应急维护方案\n读取最后一次成功的CSV快照\n人工介入引导", "仅保障最低可操作性\n需人工确认所有工单"]],
  col_widths=[1.5, 3.0, 3.5, 4.5, 4.0])

doc.add_heading('1.5 三角色权限体系', level=2)
T(["角色", "标识色", "关注焦点", "可见页面", "核心指标"],
  [["运维工程师\n(Operator)", "青色 #00BCD4", "异常告警 → 工单执行 → 根因确认", "首页(运维视图)、设备矩阵、鹰眼球体、\n仪表盘(Tab1,5,6)、AI Copilot、\n工单跟踪、工作流、报告中心", "告警精确率83.9%\n运维成本降低35%"],
   ["生产管理负责人\n(Manager)", "琥珀色 #FF9800", "KPI趋势 → 成本对比 → 策略选择 → 资源调度", "首页(管理视图)、仪表盘(Tab0,5,6,7)、\n工单跟踪、库存管理、技师管理、\n工作流、报告中心", "决策效率提升50%\n预期月度节约"],
   ["平台开发人员\n(Developer)", "紫色 #9C27B0", "算法验证 → 数据探索 → 模型迭代 → 系统扩展", "全部13个页面，含技术架构、\n知识库、仪表盘全部8Tab", "开发周期缩短60%\n算法AUC≥0.90"]],
  col_widths=[2.5, 2.5, 3.5, 5.0, 3.0])

PAGE_BREAK()

# ═══════════════════════════════════════════ 二、角色门 ═══════════════════════════════════════
doc.add_heading('二、角色门（role-gate.html）', level=1)
P("角色门是用户进入平台前看到的唯一页面。其设计目标是让用户在首次接触平台时明确自身角色，"
  "并据此获得定制化的功能视图。页面不加载导航栏、不执行角色检查——它是角色体系之前的\"零态\"页面。")

doc.add_heading('2.1 页面布局', level=2)
RP([("技术实现：", True, False, 10, None),
   ("纯静态HTML + design-tokens.css（设计令牌）+ theme-init.js（主题同步）。"
    "无Vue/React/ECharts依赖，加载极快。背景使用CSS生成的工业齿轮水印（conic-gradient旋转动画）、"
    "雷达扫描脉冲（radial-gradient动画）和15个浮动粒子（绝对定位+CSS animation随机轨迹），"
    "营造工业科技感。数据网格叠加层使用repeating-linear-gradient。", False, False, 10, None)])

CHART_BLOCK("role-cards", "三角色选择卡片", "CSS玻璃态卡片（毛玻璃效果 + 3D微旋转 + 悬停抬升动画）", "",
    "三张卡片分别展示：\n- 运维工程师(青色主题，齿轮图标)：\"异常智能告警 / 工单自动流转 / 根因快速定位\"\n- 生产管理负责人(琥珀主题，图表图标)：\"关键指标总览 / 成本效益分析 / 策略对比选择\"\n- 平台开发人员(紫色主题，扳手图标)：\"完整数据探索 / 模型快速迭代 / 系统灵活扩展\"",
    "无外部数据源——纯前端静态内容", "引导用户根据自身职责选择对应角色身份进入平台",
    "每张卡片下方的百分比指标（35%/50%/60%）为该角色的预期效率提升，基于平台内部测试估算。"
    "点击任意卡片调用selectRole(role)函数：将角色存入sessionStorage，读取URL中的?redirect=参数，"
    "跳转至目标页面（管理角色默认跳转/dashboard，其余角色跳转/）。"
    "卡片悬停时触发translateY(-8px)抬升+box-shadow放大+边框发光的三重反馈动画。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 三、首页 ═══════════════════════════════════════
doc.add_heading('三、首页 — 设备健康状态总览（home.html）', level=1)
P("首页（路由 /）是平台使用频率最高的页面。它集成了10×10设备健康矩阵、运维KPI面板、管理策略对比、"
  "SHAP可解释性面板、故障注入演示和Demo自动导览六大功能模块。以下按模块逐一详解。")

# 3.1 10x10 Grid
doc.add_heading('3.1 10×10设备健康矩阵', level=2)
P("这是首页最核心的可视化组件，以10行×10列的彩色网格同时展示全部100台CNC设备的健康状态。"
  "每个单元格代表一台设备（CNC_001至CNC_100），颜色编码反映设备的当前健康等级。")

CHART_BLOCK("machine-grid", "10×10设备健康矩阵", "纯DOM网格（100个div单元格）", "无坐标轴——网格布局",
    "100个设备单元格的颜色按健康评分五级映射：\n"
    "- 深绿色(accent-green): 健康评分≥85，设备正常运转\n"
    "- 浅绿色(#4ade80): 健康评分70-84，亚健康但无紧迫风险\n"
    "- 琥珀色(accent-amber): 健康评分55-69，存在退化迹象需关注\n"
    "- 橙色(#fb923c): 健康评分40-54，即将达到关键阈值\n"
    "- 红色(accent-red): 健康评分<40，需立即安排维护\n"
    "每个单元格右上角：有活跃工单时显示橙色角标（数字=工单数）\n"
    "P1最高优先级设备：金色呼吸脉冲动画（box-shadow glow循环1.5s）\n"
    "悬停时显示信息浮层：设备ID+健康分+故障模式+优先级徽章",
    "主数据: data/equipment_health_score.csv（health_score列决定颜色，health_status列决定标签，failure_mode列决定悬停显示）\n"
    "工单关联: data/maintenance_work_orders.csv（machine_id匹配，status='in_progress'或'pending'触发角标）\n"
    "初始加载优化: GET /api/maintenance/machines-summary（服务器预计算~15KB JSON，避免前端加载900KB z_scores.csv）",
    "运维人员一眼识别哪些设备需要关注：红色和橙色的设备需要立即或尽快安排维护。网格视图允许快速比较设备间健康差异。",
    "如何解读：点击红色/橙色设备 → 查看详情面板中的SHAP归因 → 确定哪个传感器参数是异常主因。金色脉冲的设备最优先处理（P1工单已生成，SLA倒计时中）。绿色设备健康但需关注角标——即使设备健康也可能有待执行的工单。",
    "device-grid-component.js的initDeviceGrid()函数：1. fetch CSV → PapaParse解析 → 构建machineMap(100个对象)；2. fetch工单CSV → 按machine_id分组 → 建立workOrderMap；3. renderGrid()循环100次创建div元素 → computeColor()线性插值计算颜色 → 添加事件监听。颜色映射在红绿渐变中线性插值，alpha通道随置信度浮动。")

doc.add_heading('3.2 运维视角KPI统计卡片（4张）', level=2)

STAT_CARD("ops-healthy-pct", "设备健康率",
    "百分比数值（如\"87%\"），绿色大字体",
    "health_status为'healthy'或'warning'的设备数 ÷ 100 × 100%。\n"
    "即（100 - critical_count - degrading_count）/ 100。",
    "data/equipment_health_score.csv → health_status列 → 分类计数")

STAT_CARD("ops-critical-count", "关键设备数",
    "整数值（如\"12\"），红色字体，有警戒图标",
    "health_status == 'critical' 的设备计数。\n"
    "即健康评分<40的设备数量。",
    "data/equipment_health_score.csv → health_status列 → 过滤'critical' → count")

STAT_CARD("ops-accuracy", "告警精确率",
    "固定值83.9%，紫色字体",
    "Z-Score > 2.0的告警中，后续被工单验证为真实故障的比例。\n"
    "此值在流水线中通过统计推断阶段计算：TP/(TP+FP)，\n"
    "其中TP=Z-Score告警且后续工单确认为故障，FP=Z-Score告警但无后续故障记录。",
    "alert_summary.csv → precision列\n"
    "（由stat_inference技能包在真实2999行数据上计算）")

STAT_CARD("ops-today-wos", "今日维护工单",
    "整数值（如\"8\"），青色字体",
    "maintenance_work_orders.csv中status为'in_progress'或'pending'的工单总数。",
    "data/maintenance_work_orders.csv → status列 → 过滤计数")

doc.add_heading('3.3 运维视角 — 健康分布条', level=2)
CHART_BLOCK("health-dist-bar", "健康状态分布条", "CSS flex四段比例条", "",
    "四色段从左到右依次为：\n- 绿色段（健康）：health_status='healthy'的设备占比\n- 浅绿段（亚健康）：health_status='warning'的设备占比\n- 琥珀段（退化）：health_status='degrading'的设备占比\n- 红色段（关键）：health_status='critical'的设备占比\n各段宽度按设备数量百分比动态设置",
    "data/equipment_health_score.csv → health_status列 → 分组计数 → 百分比",
    "一目了然展示100台设备的整体健康结构。",
    "如果绿色段占比>80%，整体健康良好；红色段>15%意味着需要进行系统性维护策略调整。",
    "前端JS: health_status列分组统计→各段flex-basis按百分比设置→纯CSS渲染，无图表库依赖")

doc.add_heading('3.4 运维视角 — 2×5紧急设备网格', level=2)
P("展示Top 10最高风险的设备，以2行×5列的卡片网格排列。每张卡片包含：设备ID、健康评分（大号色码数字）、"
  "距检修截止时间（倒计时格式）、诊断出的根因（如\"电压漂移 — Z_V=2.8\"）。排序依据：health_score升序，"
  "即健康评分最低的10台设备。数据来源：equipment_health_score.csv按health_score排序取前10行。")

doc.add_heading('3.5 运维视角 — 维护工单卡片', level=2)
P("从industrial_maintenance_plan.csv渲染的维护计划卡片。每张卡片以左侧色条标识优先级（P1红/P2橙/P3蓝），"
  "包含设备ID、故障模式（中文）、指定维护动作、所需备件及数量、分配技师姓名、停机窗口时间段、"
  "预计节约成本（¥）。P1工单卡片有1.5秒周期的呼吸脉冲动画。卡片数据来自当前策略（成本效率/生产效率/质量优先）"
  "下的决策生成模块输出。")

doc.add_heading('3.6 管理视角KPI面板', level=2)
P("管理角色（data-role='manager'）在首页看到4张KPI卡片：关键设备数（同运维视角但展示格式更突出成本影响）、"
  "活跃工单数（标注当前策略名称）、预期月度节约（maintenance_work_orders.csv中expected_savings列求和÷1000，"
  "单位千元）、告警精确率（同83.9%）。下方还有日成本风险Top 5表格（按cost_at_risk降序、"
  "标注风险等级和归因故障模式）和策略选择器。")

doc.add_heading('3.7 策略选择器', level=2)
P("三个策略选项：cost_efficiency（成本效率 — 最少工单数量，适合非关键生产期）、"
  "production_efficiency（生产效率 — 最少停机时间，适合旺季满产期）、"
  "quality_first（质量优先 — 最低质量风险，适合高精度加工期）。"
  "切换策略时触发POST /api/maintenance/strategy → 后端重新执行决策引擎 → "
  "生成新的工单、备件计划、停机调度CSV → 覆盖dashboard data目录 → 前端自动刷新所有相关数据显示。"
  "策略元信息卡片显示：优先目标、适应的生产阶段、预期效果指标。")

doc.add_heading('3.8 设备详情面板（右侧滑入）', level=2)
P("点击10×10网格中的任意设备单元格或2×5紧急设备卡片，触发右侧详情面板（420px宽，玻璃态背景）。"
  "面板分五部分：")
T(["面板区域", "显示内容", "数据来源", "如何解读"],
  [["① 健康指标", "设备ID、健康评分（大号色码数字）、\n退化模式（中文）、ML故障密度、\n基准评分、风险等级徽章", "equipment_health_score.csv\n对应设备行", "健康评分<40: 需立即安排维护\nML故障密度>0.7: 多信号指向异常"],
   ["② 多信号指示器", "三个子信号的状态灯和Z值：\n- Z-Score复合值（红/琥珀/绿）\n- 电压异常（|Z_V|>2.0触发）\n- 温度异常（|Z_T|>2.0触发）\n- 趋势告警（斜率>0.02触发）", "z_scores.csv\n对应设备的最新行", "信号灯状态说明哪个传感器维度触发了异常\n例如：Z_V=2.8(红) + Z_T=1.2(绿)\n→ 电压漂移问题，非热积聚"],
   ["③ SHAP归因", "四个传感器参数的SHAP水平条形图\n每个条形标注参数名和SHAP值\n条形方向（左/右）表示正向/负向贡献", "shap_dashboard.json\n对应设备的SHAP值", "哪个条形最长 → 该参数是导致异常的最主要原因\n例如：Temperature SHAP=+0.15最长\n→ 温度异常是健康评分下降的主因"],
   ["④ 维护建议", "建议的维护动作 + 所需备件 + 分配技师", "diagnosis_results.csv\n+ technician_assignments.csv", "包含具体的检修步骤和所需备件型号"],
   ["⑤ 检查清单", "4项外观检查 + 3项电气检查 + 2项备件替换\n每项可勾选", "预定义的CNC维护标准清单", "运维人员按清单执行检查并勾选确认"]])

doc.add_heading('3.9 SHAP全局探索弹窗', level=2)
CHART_BLOCK("shap-exploration", "SHAP特征贡献散点图", "ECharts散点图（bubble scatter）",
    "X轴: SHAP值（特征对模型输出的边际贡献，单位与模型输出一致）\n"
    "Y轴: 4个传感器参数（Voltage, Amperage, Temperature, Rotor Speed）\n"
    "每个点代表一台设备在一个参数上的SHAP值\n"
    "点的颜色: 该参数的特征值高低（红色=高特征值，蓝色=低特征值）",
    "4组散点（每个参数一组），共100个点/组",
    "data/shap_scatter_data.json — 100台设备×4特征的SHAP值二维数组\n"
    "由ml_inference技能包的XGBoost + SHAP TreeExplainer计算",
    "全局视角下理解每个传感器参数对健康评分的贡献方向和幅度。"
    "帮助开发人员判断：哪个传感器参数对异常检测最有信息量？",
    "看点在Y轴上的分布——如果某个参数（如Temperature）的点几乎全在0附近，说明该参数在当前模型中的贡献接近于零。\n"
    "如果某个参数的点在X轴上分布很宽（既有正SHAP也有负SHAP），说明该参数是区分正常和故障的关键特征。\n"
    "颜色模式：同一参数内，红点（高特征值）如果集中在右（正SHAP）→ 特征值高推动风险评分升高（符合物理直觉）。",
    "ml_inference技能包: XGBoost模型训练 → SHAP TreeExplainer计算每个样本×每个特征的SHAP值 → "
    "导出为shap_scatter_data.json（100×4矩阵）→ 前端ECharts scatter类型渲染")

doc.add_heading('3.10 基线追溯面板', level=2)
CHART_BLOCK("baseline-trace", "基线追溯三选项卡弹窗", "ECharts多图组合（3选项卡切换）", "",
    "三个选项卡：\n"
    "Tab 1「健康趋势」: ECharts折线图，X轴=时间步（最近30个数据点），Y轴=健康评分(0-100)。\n"
    "  绿线=健康评分趋势，红色虚线=关键阈值(40)，琥珀虚线=警告阈值(70)。\n"
    "Tab 2「Z-Score时序」: ECharts多线图，X轴=时间步，Y轴=Z-Score值。\n"
    "  4条线=Z_V(青)、Z_A(琥珀)、Z_T(红)、Z_Composite(白)，灰色虚线=±2.0阈值。\n"
    "Tab 3「成本风险分解」: ECharts堆叠柱状图+折线。\n"
    "  堆叠柱=各风险组件（故障成本/停机成本/质量成本），折线=总成本风险趋势。",
    "Tab1+Tab2: data/z_scores.csv → 按设备ID筛选 → 取最近30行\n"
    "Tab3: data/cost_risk_matrix.csv → 按设备ID筛选 → 取最近30行",
    "三选项卡提供从\"是什么\"（健康趋势）→\"为什么\"（Z-Score归因）→\"值多少\"（成本量化）的逐层深入分析。",
    "Tab1: 如果健康评分持续下降（折线左上到右下），即使仍在绿色区域，也需要预防性关注。\n"
    "Tab2: 如果Z_Composite跨越±2.0虚线并持续向上，表示多参数同步恶化。\n"
    "Tab3: 堆叠柱中如果\"故障成本\"部分占比增大，说明故障概率在上升（即使停机成本稳定）。",
    "device-grid-component.js的showTrace(type, mid)函数：\n"
    "fetch对应CSV → PapaParse解析 → 按machine_id过滤 → ECharts实例初始化 → 渲染")

doc.add_heading('3.11 故障注入演示', level=2)
P("故障注入是开发者角色的专属功能，通过模态窗口模拟向指定设备注入故障信号，演示\"信号异常→检测→诊断→处置\""
  "的完整链路。纯内存计算，不修改任何数据文件。六步动画演示：")
T(["步骤", "动画内容", "技术实现"],
  [["① 信号注入", "在内存中修改目标设备的传感器值\n（如Voltage += 偏移量×严重等级）", "POST /api/fault-injection → routes.py处理\n纯内存操作，不写文件"],
   ["② Z-Score重算", "大号Z值数字从0.5跳动到3.2\n颜色从绿→琥珀→红渐变", "CSS transition + JS setInterval数字递增"],
   ["③ SHAP诊断", "SHAP条形图更新为新计算的归因值\n显示主要贡献参数", "调用后端诊断引擎重算SHAP值"],
   ["④ 工单生成", "一张新的维护工单卡片从右侧飞入\n含设备ID/优先级/建议动作", "CSS translateX动画 + JS动态创建DOM"],
   ["⑤ 技师分配+通知", "显示分配的技师信息和邮件通知模拟", "静态演示（实际生产环境触发SMTP发送）"],
   ["⑥ 系统恢复", "所有修改回滚至原始数据，界面复位", "前端保存快照 → 关闭时恢复"]],
  col_widths=[2.0, 5.5, 5.5])

doc.add_heading('3.12 Demo自动导览模式', level=2)
P("通过URL参数 ?demo=true 激活。demo-mode.js使用sessionStorage实现跨页面状态持久化。包含8个自动推进步骤："
  "(1)10×10健康网格→(2)SHAP归因→(3)故障注入→(4)算法对比墙（自动跳转至/technical-overview）→"
  "(5)KDE数据天花板→(6)三策略模式对比（自动跳转至/dashboard）→(7)Pareto前沿3D→(8)AI Copilot（自动跳转至/chat）。"
  "界面提供聚光灯遮罩（突出当前操作区域）、解说卡片（带\"是什么/为什么/效果\"三段说明）、"
  "播放控制栏（上一步/暂停/下一步/退出按钮）、底部进度条（8步指示灯）。键盘快捷键：← → 切换步骤，Space 暂停/继续，Esc 退出。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 四、仪表盘 ═══════════════════════════════════════
doc.add_heading('四、仪表盘 — 多维度数据分析（index.html）', level=1)
P("仪表盘（路由 /dashboard）是平台功能最密集的页面，文件大小约424KB，包含8个Tab和约58个ECharts图表。"
  "以下按Tab逐一详解每个图表和卡片。")

# ── TAB 0 ──
doc.add_heading('4.1 Tab 0 — 业务总览（Manager视图）', level=2)
P("Tab 0是面向生产管理负责人的执行摘要面板，提供从100台设备到三个维护策略的全景视图。")

STAT_CARD("sec0-stats", "5张统计卡片",
    "高危设备数 | 活跃工单数 | 预期月度节省(千元) | 告警精确率(83.9%) | 传感器升级5年ROI(994%)",
    "高危设备: equipment_health_score.csv中health_level='Critical'计数\n"
    "活跃工单: maintenance_work_orders.csv中有machine_id的行数\n"
    "预期月度节省: maintenance_work_orders.csv expected_savings求和÷1000\n"
    "告警精确率: 固定值83.9%（统计推断阶段计算）\n"
    "5年ROI: 固定值994%（传感器升级成本效益分析模型计算）",
    "equipment_health_score.csv + maintenance_work_orders.csv + 统计推断阶段计算结果")

CHART_BLOCK("chart-pareto-3d", "帕累托前沿三维曲面图", "ECharts-GL 3D曲面图（WebGL加速渲染，可旋转/缩放）",
    "X轴: 总维护成本（Total Cost, $）\nY轴: 总停机时间（Total Downtime, 小时）\nZ轴: 质量风险指数（Quality Risk Index, 0-1）",
    "灰色散点: Pareto前沿面上的采样点（代表在给定成本+停机约束下可达的最优质量水平）\n"
    "绿色球体: cost_efficiency策略的锚点位置\n"
    "蓝色球体: production_efficiency策略的锚点位置\n"
    "紫色球体: quality_first策略的锚点位置",
    "data/pareto_frontier.json — 多目标优化求解得到的Pareto前沿面三维采样点\n"
    "由决策生成技能包中的多目标优化模块（NSGA-II类算法）计算",
    "展示\"不可能三角\"——成本、停机时间、质量风险三者无法同时最小化。\n"
    "三种策略在三维空间中的位置直观说明各自的取舍：\n"
    "成本效率策略靠近低成本/高停机区域，质量优先策略靠近低成本/高停机区域的对角位置。",
    "三个策略锚点在图中的空间关系是核心信息：\n"
    "如果两个策略锚点距离很近 → 这两个策略在当前数据下的实际效果差异不大\n"
    "如果某个策略锚点远离Pareto前沿面 → 该策略不是最优选择（被其他策略支配）\n"
    "鼠标拖拽旋转查看3D视角，滚轮缩放观察局部细节",
    "决策生成技能包: 以三种策略的权重参数分别运行优化 → 得到三组(成本,停机,质量)坐标 → "
    "沿Pareto前沿面采样200+点 → 导出pareto_frontier.json → 前端ECharts-GL渲染")

CHART_BLOCK("chart-strategy-comp", "三策略对比柱状图", "ECharts分组柱状图",
    "X轴: 三种策略名称（cost_efficiency / production_efficiency / quality_first）\nY轴: 双Y轴——左轴=工单数量，右轴=平均每工单成本($)",
    "蓝色柱: 工单总数（左Y轴）\n青色柱: 平均每工单成本（右Y轴）",
    "data/strategy_comparison.csv — 三种策略下的工单统计",
    "直观对比三种维护策略的资源投入（工单数=人力投入，平均成本=经济投入）。",
    "成本效率策略: 工单数最少但平均成本最高（只修最关键的，但单次维修贵）\n"
    "质量优先策略: 工单数最多但平均成本最低（预防性维护多但每次维修便宜）\n"
    "生产效率策略: 处于两者之间（平衡停机时间和成本）",
    "决策生成技能包: 分别以三种策略参数运行IndustrialMaintenanceEngine → 统计工单数和成本 → 导出CSV")

CHART_BLOCK("chart-pareto-marginal", "边际成本路径分析", "ECharts连线散点图（带标注的折线+散点）",
    "X轴: 累计停机时间减少量（小时）\nY轴: 累计额外成本（$）",
    "折线连接三个策略切换节点的边际成本轨迹\n每个节点标注策略名称和边际成本率($/hr)",
    "data/pareto_frontier.json + 策略对比数据",
    "展示每减少一小时停机时间需要付出的额外成本——帮助管理层判断\"为减少停机花多少钱是值得的\"。",
    "折线斜率=边际成本率。斜率陡峭时→减少停机非常昂贵；斜率平缓时→减少停机成本效益好。\n"
    "在斜率突变点标注策略切换，说明从哪个点开始边际效益递减。")

PAGE_BREAK()

# ── TAB 1 ──
doc.add_heading('4.2 Tab 1 — 运行日志（Operator视图）', level=2)
P("Tab 1为运维工程师提供设备传感器时序数据的查询和可视化。核心功能是按设备查看历史传感器读数和故障分布。")

CHART_BLOCK("chart-log-ts", "参数时序折线图", "ECharts多线折线图",
    "X轴: 时间戳（datetime，来自log.csv的Date列，范围约30天）\nY轴: 传感器原始值（各自物理单位：V/A/°C/RPM）",
    "4条折线（不同颜色）：\n- Voltage（青色 #66d9c8）：电压时序\n- Amperage（琥珀色 #ffb340）：电流时序\n- Temperature（红色 #ff453a）：温度时序\n- Rotor Speed（蓝色 #6db5f9）：转子转速时序",
    "data/log.csv — 按选定设备ID过滤 → Date, Op.Voltage, Op.Amperage, Op.Temperature, Rotor Speed列",
    "运维人员选择一台设备后，查看其四种传感器参数在过去30天内的变化趋势，判断是否有异常波动。",
    "正常: 四条线各自在稳定范围内小幅波动（无剧烈跳变）\n"
    "电压漂移: Voltage线出现持续上升或下降趋势（漂移>基线标准差的2倍）\n"
    "热积聚: Temperature线持续上升且超过基线均值+2σ\n"
    "功率异常: Voltage和Amperage同时出现异常波动（相关性突变）\n"
    "点击图表中的数据点可标记查看具体数值",
    "前端: PapaParse加载log.csv → 按设备ID过滤 → ECharts dataset + 4个series → 渲染")

CHART_BLOCK("chart-log-fault", "故障类型分布饼图", "ECharts饼图",
    "环形饼图，每扇区代表一种故障类型的记录占比",
    "10种故障类型（Type 0-9）各占一个扇区\n颜色: 10色调色板（红→橙→黄→绿→青→蓝→紫→粉→棕→灰）\n扇区大小=该设备上该故障类型的记录数÷总记录数",
    "data/log.csv → 按设备ID过滤 → Failure.Equipment.Type列 → 分组计数",
    "展示选定设备的历史故障记录中，各故障类型的分布比例。帮助判断该设备的主要故障模式。",
    "如果某个故障类型的扇区占比>50% → 该设备有明确的\"主导故障模式\"\n"
    "如果扇区分布均匀 → 该设备的故障模式多样化，需要综合诊断\n"
    "Type 0占比大可能与误报或正常运行标签相关（需结合领域知识判断）")

P("Tab 1还包含一个可排序、可分页（每页50条）的原始数据表格，列包括Date/Equipment.Id/Failure.Type/"
  "Voltage/Amperage/Temperature/Rotor Speed。支持按设备ID搜索、故障类型chip筛选、电压范围过滤。"
  "数据来源：data/log.csv（2999行原始传感器数据）。")

PAGE_BREAK()

# ── TAB 2 ──
doc.add_heading('4.3 Tab 2 — 数据探索（Developer视图）', level=2)
P("Tab 2面向开发人员，提供从原始数据到跨表关联分析的完整探索工具链。包含5个ECharts图表+20张分析图片。")

CHART_BLOCK("chart-fault-stack", "故障类型堆叠柱状图", "ECharts堆叠柱状图",
    "X轴: 100台设备ID（标签旋转90度）\nY轴: 故障记录条数",
    "10个堆叠系列=10种故障类型(Type 0-9)\n每柱总高度=该设备的总故障记录数\n各色段高度=该设备每种故障的记录数",
    "data/fault_dist.csv — 由data_prep技能包从log.csv按设备+故障类型交叉统计生成",
    "全局视角下100台设备的故障类型构成。识别\"问题设备\"（总记录过多）和\"问题类型\"（某类型普遍存在）。",
    "柱体极高的设备→故障记录异常多，需要审查是否设备本身有问题还是数据采集异常\n"
    "某颜色在所有设备中都有较大占比→该故障类型是全产线共性问题\n"
    "颜色集中在少��设备→局部问题，可能是特定批次或工况导致")

CHART_BLOCK("chart-boxplot", "参数箱线图：正常 vs 故障", "ECharts箱线图（boxplot）",
    "X轴: 8个类别=V正常, V故障, A正常, A故障, T正常, T故障, RPM正常, RPM故障\nY轴: 传感器原始值（各自物理单位）",
    "2个系列：\n- 绿色箱体: 正常运行记录的分布（min-Q1-median-Q3-max）\n- 红色箱体: 故障记录的分布（min-Q1-median-Q3-max）",
    "data/log.csv → 按Failure.Equipment.Type=0划分正常，≠0划分故障 → 计算每个参数的箱线统计量(min/Q1/median/Q3/max)",
    "直观对比正常和故障状态下各传感器参数的分布差异。如果绿色和红色箱体高度重叠→该参数单独无法区分正常和故障。",
    "箱体重叠度=该参数的判别能力。重叠越多→Youden's J越低。\n"
    "中位数偏差=Median(故障)-Median(正常)，偏差越大越有区分力。\n"
    "此图直接可视化论证\"4参数条件下纯ML检测能力有限\"——4对箱线图中至少有2-3对高度重叠。")

CHART_BLOCK("chart-fault-scatter", "设备故障负载散点图", "ECharts气泡散点图",
    "X轴: 故障记录条数\nY轴: 不同故障类型数量",
    "每个气泡=一台设备\n气泡大小: 正常记录条数（越大=正常运行时间越长）\n气泡颜色: 故障率（绿=低, 琥珀=中, 红=高）\n标签: 设备ID后两位数字",
    "data/fault_dist.csv — 按设备统计故障记录数和故障类型数",
    "识别\"问题设备\"：哪些设备故障多且故障类型多样？右上角的气泡最需要关注。",
    "右上角区域(高X+高Y)=故障频繁且类型多样的\"问题设备\"\n"
    "右下角区域(高X+低Y)=单一故障类型反复出现→可能是系统性缺陷\n"
    "气泡大+颜色绿=设备总体运行时间长且故障率低→\"健康设备\"\n"
    "气泡小+颜色红=运行时间短但故障率高→\"新设备异常\"或\"数据不完整\"")

CHART_BLOCK("chart-cross-corr", "跨表关联矩阵热力图", "ECharts热力图（heatmap）",
    "X/Y轴: 6个跨表维度=故障率, 参数CV%, 日产值, 成本风险, 距维修天数, 距保养天数\n单元格: Spearman秩相关系数ρ",
    "颜色从红(-1, 负相关) → 白(0, 无关) → 绿(+1, 正相关)\n每格标注ρ值",
    "data/cross_table_metrics.csv — 从log.csv(故障数据)+summary.csv(生产数据)+maintenance数据联合计算",
    "发现不同数据表（LOG/SUMMARY/ASSEMBLY/TESTS）中指标之间的隐藏关联。",
    "强正相关(深绿, ρ>0.5): 两个指标同向变动→可能是一个驱动另一个\n"
    "强负相关(深红, ρ<-0.5): 两个指标反向变动→存在tradeoff\n"
    "接近白色(|ρ|<0.1): 两个指标基本独立\n"
    "例如: \"故障率\"与\"参数CV%\"正相关→参数越不稳定(高CV)，故障率越高(符合物理直觉)")

CHART_BLOCK("chart-cross-heatmap", "故障组×保养周期成本风险热力图", "ECharts热力图",
    "X轴: 保养周期间隔分类\nY轴: 故障组分类\n单元格值: 平均成本风险($)",
    "颜色从绿(低成本风险) → 红(高成本风险)\n每格标注平均成本风险值",
    "data/cross_table_metrics.csv — 按故障组×保养周期交叉分组计算均值",
    "识别哪种故障类型+保养周期组合的成本风险最高，指导保养策略优化。",
    "红色单元格→该故障组在该保养周期下的平均成本风险极高→需要缩短保养周期\n"
    "绿色单元格→当前保养周期对该故障组是充分的\n"
    "此行热力图直接回答:\"如果我把某故障组的保养周期从X改为Y，能降低多少成本风险？\"")

CHART_BLOCK("chart-cross-r3r4", "参数CV% × 超规格率散点图", "ECharts散点图",
    "X轴: 参数变异系数CV%均值（传感器不稳定度）\nY轴: 产品超规格率%（质量缺陷率）",
    "每个点=一台设备\n青色圆点\n数据点上方标注设备ID",
    "data/cross_table_metrics.csv — 15台有完整四表数据的设备",
    "验证\"传感器不稳定→产品质量下降\"这一因果假设。",
    "如果散点从左下到右上呈对角分布→传感器CV%与超规格率正相关(不稳定确实影响质量)\n"
    "如果散点随机分布→两者无关，质量问题另有原因\n"
    "右上角的点=传感器不稳定且质量差→优先升级该设备的传感器")

CHART_BLOCK("chart-chain1-sankey", "传导链：保养周期→参数稳定性→超规格率", "ECharts桑基图（sankey diagram）",
    "三列流向图：保养周期分组 → 参数CV%稳定性分组 → 超规格率分组",
    "节点: 保养周期(长/中/短) → CV%(稳定/一般/不稳定) → 超规格率(低/中/高)\n边宽=条件概率(标注Wilson 95%置信区间)",
    "data/chain1_conditional_prob.csv — 条件概率矩阵",
    "展示\"保养频率→传感器健康→产品质量\"的传导路径和量化强度。",
    "粗流量路径=高概率传导链。例如：短保养周期→CV%稳定→低超规格率(粗流)=频繁保养确实保障了质量。\n"
    "如果存在\"短保养周期→高超规格率\"的细流→这些设备的故障不是保养能解决的(需要更换而非保养)。\n"
    "Wilson 95%CI标注了概率估计的可靠性——区间宽=样本少估计不可靠")

CHART_BLOCK("chart-chain2-mediation", "中介效应Bootstrap分析", "ECharts水平柱状图",
    "Y轴: 四种效应=总效应(Total Effect), 间接效应(Indirect a*b), 直接效应(Direct Effect c'), 控制变量\nX轴: 效应值",
    "水平柱=效应估计值（bootstrap均值）\n误差线=95% BCa Bootstrap置信区间\n灰色竖虚线: x=0（零效应参考线）",
    "data/chain2_mediation_bootstrap.csv — 1000次Bootstrap重采样结果",
    "检验\"故障类型→成本风险→保养紧迫度\"这条中介路径是否成立。"
    "如果间接效应的CI不包含0→成本风险是显著中介变量；如果包含0→成本风险不是中介。",
    "间接效应(a*b)的95%CI如果完全在0的右侧(不跨越灰色虚线)→存在正向中介效应(显著)\n"
    "如果CI跨越0→中介效应不显著（当前数据不支持\"成本风险中介\"的假设）\n"
    "此图结果: 三条路径的CI均跨越0→\"成本风险中介故障类型对保养紧迫度的影响\"这一假设在当前数据中不成立。"
    "这意味着需要寻找其他中介变量来解释故障类型如何影响保养紧迫度。")

PAGE_BREAK()

# ── TAB 3 ──
doc.add_heading('4.4 Tab 3 — 基线划定分析（Developer视图）', level=2)

CHART_BLOCK("chart-variance", "方差分解：设备间 vs 设备内", "ECharts堆叠柱状图",
    "X轴: 4个传感器参数（Voltage, Amperage, Temperature, Rotor Speed）\nY轴: 方差占比(%)",
    "两个堆叠层：\n- 深色层（设备间方差Between-device）: 不同设备之间的均值差异导致的方差\n- 浅色层（设备内方差Within-device）: 同一设备在不同时间的波动导致的方差\n每柱总高=100%",
    "data/variance_decomposition.csv — 单因素ANOVA方差分解结果\n由stat_inference技能包在100台设备×30天数据上计算",
    "论证\"逐设备基线是强制性的\"——设备间方差占比61-73%意味着不同设备的正常运转范围本就不同，"
    "不能用统一的全局阈值来判断所有设备是否异常。",
    "如果设备间方差>60%（如此图所示）→必须为每台设备建立独立基线\n"
    "如果设备间方差<20%→可以使用全局阈值（但本数据不满足此条件）\n"
    "设置不同基线后，\"异常\"的定义从\"偏离全局均值\"变为\"偏离该设备自身历史正常范围\"——精度大幅提升")

CHART_BLOCK("chart-alert-pie", "告警等级分布饼图", "ECharts饼图",
    "环形饼图，每扇区代表一种告警等级的设备占比",
    "4个扇区：\n- ALARM（红色）: Z_comp>2.5\n- WARNING（琥珀色）: Z_comp>2.0\n- WATCH（蓝色）: Z_comp>1.5\n- NORMAL（绿色）: Z_comp≤1.5",
    "data/alert_summary.csv — 100台设备的最新Z-Score分级统计",
    "全局告警态势——当前有多少设备处于各告警等级。",
    "ALARM+WARNING占比>20%→系统整体健康状况需要关注\n"
    "NORMAL占比>70%→多数设备运转正常（但需注意False Negative风险——Z-Score≤1.5也不代表绝对正常）")

CHART_BLOCK("chart-baseline-src", "基线来源分布", "ECharts柱状图",
    "X轴: 三种基线来源\nY轴: 使用该基线来源的设备数量",
    "三种来源:\n- 自基线(Self): 该设备有≥3个历史样本→使用自身基线\n- 混合基线(Mixed): 自身样本不足→使用同型号均值基线\n- 集群回退(Cluster): 同型号也样本不足→使用全局中位数基线",
    "data/baseline_stats.csv — baseline_source列统计",
    "展示基线建立的三层回退策略的实际使用分布。理想情况下大部分设备应使用自基线。",
    "自基线占比越高→数据越充分→基线越准确\n"
    "集群回退占比高→可能需要增加数据采集频率或延长采集周期\n"
    "如果集群回退设备恰好与告警设备重合→需要额外关注(基线不准确可能导致误判)")

CHART_BLOCK("chart-zscore-ts", "Z-Score时序图（单设备）", "ECharts多线折线图",
    "X轴: 时间步（最近30个数据点）\nY轴: Z-Score值",
    "4条折线: Z_Voltage(青), Z_Amperage(琥珀), Z_Temperature(红), Z_Composite(白)\n灰色虚线: Z=±2.0 (WARNING阈值)\n红色虚线: Z=±2.5 (ALARM阈值)",
    "data/z_scores.csv — 用户从下拉菜单选择设备 → 过滤该设备的最近30行",
    "查看单台设备各参数Z-Score的历史演变。Z-Score>2.0意味着该参数的当前值偏离其历史均值的幅度超过了2个标准差。",
    "正常: 各参数Z-Score在±2.0范围内上下波动（白噪声特征）\n"
    "异常趋势: Z值持续向上→参数在持续恶化(不是偶然波动)\n"
    "跳变: Z值突然从0.5跳到3.0→可能发生了突发故障\n"
    "如果Z_Composite(白线)和Z_Temperature(红线)同步上升→热积聚问题")

CHART_BLOCK("chart-radar", "故障类型参数签名雷达图", "ECharts雷达图",
    "3-4轴: 每个故障类型的参数Z-Score特征(Voltage/Amperage/Temperature/Rotor Speed)",
    "每条折线=一种故障类型的参数签名\n各轴上的点位=该故障类型下该参数的平均Z-Score",
    "data/failure_sig.csv — 按故障类型分组计算各参数的平均Z-Score",
    "展示每种故障类型的\"参数指纹\"——即在哪种故障下、哪个传感器参数会表现出多大的异常。",
    "例如: 如果Type 1故障在Voltage轴上伸出很远(高Z_V)但在Temperature轴上接近中心→"
    "Type 1是电压相关故障\n"
    "如果两种故障类型的雷达图形状相似→它们在参数空间中的表现相似，可能难以仅凭4参数区分\n"
    "每种故障的\"签名\"可用于快速诊断: 新出现的异常设备匹配哪种故障的签名→初步判断故障类型")

CHART_BLOCK("chart-bubble", "成本风险矩阵气泡图", "ECharts气泡图",
    "X轴: 故障概率（0-1）\nY轴: 单位产值（$/天）",
    "每个气泡=一台设备\n气泡大小: 日产量(件)\n气泡颜色: 健康评分（绿=高→红=低）",
    "data/cost_risk.csv — cost_risk_matrix.csv的前端可视化",
    "识别\"高风险×高价值\"设备——即故障概率高且单位产值高的设备，这些设备是维护优先级最高的。",
    "右上角(高故障概率+高产值): 最需要立即维护的设备——一旦故障将造成最大经济损失\n"
    "大尺寸(高日产量)+右上角=绝对最优先级的维护目标\n"
    "颜色越红=健康评分越低=越紧迫\n"
    "此图的业务语言:\"如果CNC_036停一天产，损失将超过$7,000\"——管理层可以直接理解")

PAGE_BREAK()

# ── TAB 4 ──
doc.add_heading('4.5 Tab 4 — 预测性维护模型（Developer视图）', level=2)

CHART_BLOCK("chart-algo-compare", "7种算法对比实验", "ECharts分组柱状图",
    "X轴: 7种算法名称\nY轴: 指标值(%)",
    "5个指标分组（每组7根柱）:\nAUC, Accuracy, Precision, Recall, F1\n每种算法有5根柱（不同颜色）",
    "data/algorithm_comparison.csv — benchmark_algorithms.py在真实2999行数据上的10折CV结果",
    "在4参数条件下公平对比主流ML算法的异常检测能力，验证\"纯ML方法AUC上限≈0.537\"的结论。",
    "所有算法的AUC都在0.50-0.59之间→任何单模型都无法在4参数数据上取得可用性能\n"
    "MTNN（多任务神经网络）通常AUC最高(≈0.59)但仍远低于0.70的可用阈值\n"
    "Recall普遍低于Precision→模型倾向于保守预测(宁可漏报也不错报)\n"
    "核心结论: 此图为\"多信号融合是必要而非可选的\"提供数据支撑",
    "benchmark_algorithms.py: 读取2999行原始数据 → 4参数特征 → 10折CV训练 → "
    "7种算法逐个评估 → 导出algorithm_comparison.csv。强调: 使用真实训练数据，非合成数据。")

CHART_BLOCK("chart-eval-bar", "XGBoost窗口大小对比", "ECharts分组柱状图",
    "X轴: 4种窗口大小（5/10/15/20个时间步）\nY轴: 指标值",
    "5个指标: Accuracy, Precision, Recall, F1, AUC\n每个窗口有5根柱",
    "data/eval_metrics.csv — 不同窗口大小下的XGBoost评估结果",
    "评估输入时间窗口大小对模型性能的影响，选择最优窗口。",
    "最优窗口通常是Recall和F1同时最高的那个\n窗口太小→信息不足；窗口太大→引入噪声\n"
    "本数据中窗口10-15通常最优")

CHART_BLOCK("chart-feat-imp", "Top 15特征重要性", "ECharts水平柱状图",
    "Y轴: 15个最重要的特征名称\nX轴: 特征重要性得分（XGBoost gain）",
    "水平柱=重要性，按降序排列",
    "data/feature_imp.csv — XGBoost训练后导出的feature_importances_",
    "识别哪些特征对模型预测最关键。",
    "排名前几的特征是模型\"最关注\"的信号，这些特征的质量直接决定了模型上限。"
    "如果排名靠前的特征都是某单一类型（如全是Z-Score衍生特征），说明模型过于依赖单一信号源。")

CHART_BLOCK("chart-robustness", "鲁棒性测试热力图", "ECharts热力图",
    "X轴: 特征组\nY轴: 噪声水平\n单元格: 性能保留率",
    "颜色从红（性能大幅下降）→绿（性能保持良好）",
    "data/robustness.csv — 向不同特征组注入不同比例噪声后重新评估的F1保留率",
    "评估模型对数据噪声和特征缺失的容忍度。",
    "某特征组在低噪声水平就变红→模型严重依赖该特征组→该特征组若不可用则模型失效\n"
    "全绿→模型对噪声鲁棒（理想但罕见）\n此图直接为\"多信号融合\"提供反面论证:如果只依赖单一模型,特征缺失将导致性能崩溃")

CHART_BLOCK("chart-variant", "MTNN变体对比", "ECharts分组柱状图",
    "X轴: MTNN架构变体名称\nY轴: 指标值",
    "不同神经网络架构变体的性能对比",
    "data/variant_comp.csv",
    "对比不同神经网络架构选择的性能差异。",
    "最优变体即为后续分析中采用的ML模型配置。")

PAGE_BREAK()

# ── TAB 5 ──
doc.add_heading('4.6 Tab 5 — 方案有效性验证（全部角色）', level=2)
P("Tab 5通过时间回溯验证（Backtesting）评估系统在实际历史数据上的检测表现。这是论证\"系统有用\"的关键证据。")

STAT_CARD("backtest-stats", "6张回溯指标卡片",
    "平均预警提前量(分钟) | 中位预警提前量(分钟) | 漏报率(%) | 检出率(%) | 总故障事件数 | 步进收敛步数",
    "在历史数据上模拟检测流程：用前N天数据建立基线 → 预测后续M天的故障 → 与实际故障记录比对\n"
    "THRESHOLD WATCH: Z>1.5 → 检出率约70%, 平均提前约40分钟\n"
    "THRESHOLD WARNING: Z>2.0 → 检出率约60%, 平均提前约60分钟(但漏报率升高)\n"
    "THRESHOLD ALARM: Z>2.5 → 检出率约35%, 平均提前约90分钟(但漏报率极高)",
    "data/backtest_lead_time_summary.csv + data/backtest_by_fault_group.csv\n"
    "由stat_inference技能包的回溯验证模块在历史数据上计算")

CHART_BLOCK("chart-backtest-leadtime", "预警提前量分布", "ECharts直方图/柱状图",
    "X轴: 预警提前量区间（分钟段）\nY轴: 检出事件数",
    "每个柱=在该提前量区间内被成功检出的事件数\n三种阈值（Watch/Warning/Alarm）可切换显示",
    "data/backtest_lead_time_summary.csv — 按阈值版本分文件(Watch/Warning/Alarm)",
    "展示系统在实际历史故障发生前多久能发出预警。提前时间越长→运维人员准备时间越充分。",
    "柱体集中在右侧(长提前量)→系统能在故障发生前充分预警\n"
    "柱体集中在左侧(短提前量)→系统只能紧急告警(留给运维的反应时间不足)\n"
    "对比三种阈值: WATCH检出的多但提前量短, ALARM提前量长但检出少→需要权衡")

CHART_BLOCK("chart-backtest-faultgroup", "故障组检出率对比", "ECharts分组柱状图",
    "X轴: 故障组名称\nY轴: 检出率(%)",
    "每种故障类型的检出率柱\n三种阈值版本可切换",
    "data/backtest_by_fault_group.csv",
    "展示不同故障类型的检出难度差异——有些故障容易被系统检测，有些则难以检测。",
    "检出率>80%的故障组→系统对该类型故障检测可靠\n"
    "检出率<30%的故障组→4参数条件下该类型故障难以检测→需要额外传感器或人工巡检补充\n"
    "此图直接为\"传感器升级优先级\"提供依据: 先为检出率最低的故障类型增加传感器维度")

CHART_BLOCK("chart-youdens", "Youden's J单参数判别力", "ECharts柱状图",
    "X轴: 4个传感器参数（Voltage, Amperage, Temperature, Rotor Speed）\nY轴: Youden's J值",
    "4根柱=4个参数各自的Youden's J\n红色虚线: J=0.30（\"可用\"阈值）",
    "从log.csv的2999行数据中，对每个参数独立计算最佳分类阈值下的(TPR, FPR) → J=TPR-FPR",
    "展示每个传感器参数单独用于区分正常/故障的能力上限。所有4个参数的J值均远低于0.30→任何一个单一传感器都不具备可靠的故障检测能力。",
    "所有柱体都在0.30红线以下→论证\"单传感器不可用\"\n"
    "柱体高度: Temperature(0.065)≈Amperage(0.064)>Rotor(0.061)>Voltage(0.077)\n"
    "注意: 即使是最高的Voltage(0.077)也仅为可用阈值的约1/4",
    "前端JS: 遍历4个参数，使用log.csv中的Failure.Type列作为标签 → 对每个参数计算ROC → 取max(TPR-FPR) → 渲染")

PAGE_BREAK()

# ── TAB 6 ──
doc.add_heading('4.7 Tab 6 — 智能维护决策中心（全部角色）', level=2)
P("Tab 6是决策中枢，展示流水线最终产出：维护工单、技师分配、备件需求、停机调度。包含约18个图表和交互组件。")

STAT_CARD("sec6-roi-stats", "ROI统计行（4张卡片）",
    "预期节省总额($) | 预防性维护总成本($) | ROI(节省÷投入) | 技术员降级节省($)",
    "从maintenance_work_orders.csv统计：\n- 节省总额=SUM(expected_savings)\n- 维护成本=SUM(estimated_cost)\n- ROI=节省/投入\n- 降级节省=使用初级技师替代高级技师的成本节约",
    "data/maintenance_work_orders.csv")

CHART_BLOCK("chart-roi-efficiency", "工单ROI排序（节省/投入比）", "ECharts水平柱状图",
    "Y轴: 设备ID（按ROI降序排列）\nX轴: 节省÷投入比值",
    "每根水平柱=一台设备的工单ROI\n颜色: 由ROI值映射（绿=高ROI, 红=低ROI）",
    "maintenance_work_orders.csv → 计算每工单的expected_savings/estimated_cost → 排序",
    "帮助管理层识别\"投入产出比最高\"的维护行动——ROI最高的设备优先分配资源。",
    "ROI>3: 每投入1元节省3元以上→高优先级\nROI<1: 投入大于节省→需要审查该工单是否必要或是否有更经济的方案")

CHART_BLOCK("chart-pareto-3d-sec6", "Pareto前沿3D散点图（决策版）", "ECharts-GL 3D散点图（可旋转/缩放）",
    "X轴: 总成本($)\nY轴: 总停机(小时)\nZ轴: 质量风险指数",
    "灰色点: Pareto前沿采样点\n3个大球: 三种策略锚点(绿/蓝/紫)",
    "data/pareto_frontier.json",
    "同Tab0的3D曲面，此处放在决策上下文中供策略对比参考。",
    "三个策略锚点在图中的空间关系是核心信息。如果两个策略锚点距离很近→这两个策略在当前数据下的实际效果差异不大。如果某个策略锚点远离前沿面→该策略被其他策略支配。")

CHART_BLOCK("chart-strategy-comp-sec6", "三策略工单数量与平均成本对比", "ECharts分组柱状图",
    "X轴: 三种策略\nY轴: 工单数量/平均成本",
    "每组两根柱: 工单数(蓝)+平均成本(青)",
    "data/strategy_comparison.csv",
    "对比三种策略的资源消耗。",
    "策略选择=在工单数量(人力投入)和平均成本(经济投入)之间取舍。")

CHART_BLOCK("chart-anomaly-donut", "异常模式分布环形图", "ECharts环形图（donut, 内径60%）",
    "环形，每扇区=一种异常模式",
    "4个扇区:\n- voltage_drift(电压漂移, 青)\n- thermal_buildup(热积聚, 红)\n- power_anomaly(功率异常, 琥珀)\n- combined_degradation(复合退化, 紫)",
    "data/diagnosis.csv — diagnosis技能包对100台设备的根因分类结果",
    "哪类异常模式在当前的100台设备中占主导？",
    "某扇区占比>50%→当前产线的主要问题是该类异常\n"
    "例如: thermal_buildup占比最高→可能需要检查冷却系统或环境温控\n"
    "combined_degradation占比高→多数设备的多个传感器参数同时恶化→可能是系统性老化")

CHART_BLOCK("chart-sched-gantt", "14天滚动维护排程甘特图", "ECharts自定义甘特图",
    "X轴: 日期（D1-D14，14天滚动窗口）\nY轴: 维护任务行",
    "每个横条=一个维护任务\n条长=预计工时\n颜色=技师类型（青=电气, 琥珀=热控, 紫=高级, 灰=初级）",
    "downtime_schedule.csv + technician_schedule.csv\n由决策生成技能包的贪心+2-opt局部搜索调度算法计算",
    "为管理层提供未来14天的维护排程可视化。确保没有技师超负荷安排、没有关键设备扎堆停机。",
    "同一天内如果横条堆积过多→该天维护资源紧张→需要调整排程\n"
    "同一设备在14天内出现多次→可能需要延长单次维护时间(治标不治本)\n"
    "某类技师（如电气专员）的横条在连续几天都出现→该技师可能需要减负或增援",
    "调度算法: 贪心初始化(按优先级+deadline排序分配) → 2-opt局部搜索(交换相邻任务减少冲突) → 输出每任务的设备/时间/技师")

CHART_BLOCK("chart-sched-load", "每日负载分布（通道容量利用率）", "ECharts堆叠柱状图",
    "X轴: 日期（D1-D14）\nY轴: 所需工时（小时）",
    "堆叠层=不同技师类型的工作时间\n虚线=总可用工时容量",
    "technician_schedule.csv — 按日汇总各技师类型的分配工时",
    "检验维护排程中是否存在某天的总工时超过技师总容量。",
    "柱体超过虚线→该天超负荷→需调整(提前/推迟/增加技师)\n"
    "柱体远低于虚线→该天有空闲→可以安排额外任务\n"
    "连续几天超负荷→需考虑招人/加班/外包")

CHART_BLOCK("chart-or-risk-compare", "运筹优化: 贪心 vs 背包优化风险削减对比", "ECharts水平柱状图",
    "Y轴: 设备ID（按风险削减差异排序）\nX轴: 风险削减量",
    "橙色柱: 贪心算法的风险削减\n紫色柱: 0-1背包DP最优解的风险削减",
    "maintenance_work_orders.csv → 每工单的成本和风险削减估计 → DP求解",
    "展示运筹优化（背包算法）相比简单贪心规则在给定预算约束下能多削减多少风险。",
    "紫色柱长于橙色柱的差距=优化带来的额外风险削减\n"
    "差距大的设备→贪心规则低估了该设备（可能是成本适中但风险极高的设备）\n"
    "差距小的设备→贪心和优化选择相同（直观规则已经足够好）",
    "0-1背包DP: 每工单=一个物品(成本=estimated_cost, 价值=expected_risk_reduction) → "
    "DP求解max(总价值) s.t. 总成本≤预算 → 对比贪心(按价值/成本比排序)")

CHART_BLOCK("chart-or-before-after", "运筹优化 vs 规则驱动四维效果对比", "ECharts分组柱状图",
    "X轴: 4个维度=风险削减, 成本效率, 资源利用率, 覆盖率\nY轴: 评分值",
    "蓝色柱: 规则驱动（优化前）\n紫色柱: 运筹优化（优化后）",
    "从贪心vs DP的结果中提取四维指标对比",
    "量化运筹优化相比规则驱动的提升幅度。",
    "紫柱高于蓝柱的维度→运筹优化带来了明确改善\n"
    "某维度无差异→该维度规则已经最优\n"
    "此图是\"运筹优化模块是否有价值\"的直接检验")

PAGE_BREAK()

# ── TAB 7 ──
doc.add_heading('4.8 Tab 7 — 设备健康与感知升级（Manager视图）', level=2)
P("Tab 7面向管理层，将设备健康评分与传感器升级ROI分析结合，为\"是否投资增加传感器\"的决策提供数据支撑。")

CHART_BLOCK("chart-health-rul-quadrant", "健康-RUL象限图", "ECharts散点图",
    "X轴: 健康评分（0-100）\nY轴: 剩余可用寿命RUL（天）",
    "每个点=一台设备\n四条象限分割线: X=60(健康分界线), Y=60(寿命分界线)\n"
    "颜色编码: 左上(高RUL+低健康)=退化中但寿命还长(琥珀), "
    "左下(低RUL+低健康)=急需维护(红), "
    "右上(高RUL+高健康)=良好(绿), "
    "右下(低RUL+高健康)=寿命将尽但当前健康(蓝)",
    "data/equipment_health_score.csv(health_score列) + data/rul_degradation.csv(RUL列)",
    "将设备健康状态和剩余寿命两个维度放在一起观察，识别不同处置策略的设备群。",
    "左上象限(退化但寿命长): 应安排预防性维护——不是最紧迫但需要计划\n"
    "左下象限(退化且寿命短, 红区): 第一优先级——紧急维护\n"
    "右下象限(健康但寿命短): 设备即将到设计寿命但当前状况良好——安排更换计划而非紧急维修\n"
    "右上象限(健康且寿命长, 绿区): 无需操作——继续按计划巡检",
    "健康评分=stat_inference计算; RUL=基于历史退化曲线+当前健康评分的Weibull回归估计")

CHART_BLOCK("chart-health-dist-pie", "健康状态分布饼图", "ECharts环形饼图",
    "环形，每扇区=一种健康等级",
    "4个扇区:\n- 健康 Healthy(绿): health_score≥85\n- 亚健康 Warning(浅绿): 70-84\n- 退化 Degrading(琥珀): 55-69\n- 关键 Critical(红): <40",
    "data/equipment_health_score.csv → health_status列分类统计",
    "全局设备健康分布概览——管理层最关心的\"我们的设备总体状态如何\"。")

CHART_BLOCK("chart-phase-timeline", "传感器升级三阶段路线图", "ECharts时间线柱状图",
    "X轴: 3个升级阶段\nY轴: 投资额($)",
    "每阶段柱体+ROI标签+Youden's J标注",
    "data/sensor_upgrade_plan.csv — 基于KDE分析模型的升级方案",
    "三阶段传感器升级的投资路线: 阶段1(振动传感器,$50K)→阶段2(电流频谱,$120K)→阶段3(红外热像,$200K)\n"
    "每阶段标注新增Youden's J增量、5年ROI和回本周期。")

CHART_BLOCK("chart-youdens-curve", "Youden's J累积提升曲线", "ECharts折线图",
    "X轴: 传感器升级阶段（Phase 0-3）\nY轴: Youden's J值（0-1）",
    "折线: Youden's J从0.075(Phase 0)→0.35(Phase 1)→0.65(Phase 2)→0.90+(Phase 3)\n红色虚线: J=0.30(可用阈值)",
    "data/sensor_phase_summary.csv",
    "直观展示每增加一种传感器维度后，系统对正常/故障的区分能力提升了多少。",
    "折线跨越0.30红线=从\"不可用\"到\"可用\"的质变\nPhase 0到Phase 1的斜率最大→振动传感器是性价比最高的第一笔投资\n"
    "Phase 3趋近1.0→7参数条件下几乎可以完美区分正常和故障(但投资也最大)")

CHART_BLOCK("chart-sensor-roi", "ROI分析三场景概率模型（5年）", "ECharts柱状+折线组合图",
    "X轴: 3个升级阶段\nY轴: 5年ROI(%)",
    "3种场景的柱:\n- 保守(Conservative, 灰): 故障率改善20%\n- 基准(Base, 蓝): 故障率改善35%\n- 乐观(Optimistic, 绿): 故障率改善50%\n折线=投资回收期(月)",
    "data/sensor_roi_analysis.csv — 蒙特卡洛模拟(1000次)的中位数和P25/P75",
    "在乐观/基准/保守三种假设下，量化每个升级阶段的投资回报。帮助管理层在不确定性下做决策。",
    "即使最保守假设下ROI仍>100%→投资几乎确定有正回报\n"
    "如果保守假设ROI<0→投资有风险，需审慎\n回收期<12个月→值得快速推进")

CHART_BLOCK("chart-capex-opex", "CAPEX vs OPEX成本结构", "ECharts堆叠柱状图",
    "X轴: 3个升级阶段\nY轴: 成本($)",
    "堆叠层:\n- 深色底层=CAPEX(一次性硬件采购+安装)\n- 浅色顶层=OPEX(年度维护+校准+软件许可)",
    "data/sensor_upgrade_plan.csv",
    "将总成本分解为一次性资本支出和经常性运营支出。",
    "CAPEX占比大→一次性投资高但后续成本低(适合有预算时一次性投入)\n"
    "OPEX占比大→持续支出模式(适合分期投入)")

PAGE_BREAK()

# ═══════════════════════════════════════════ 五、AI Copilot ═══════════════════════════════════
doc.add_heading('五、AI Copilot — 智能对话助手（chat.html）', level=1)
P("AI Copilot的独特价值在于：它是\"知道平台实时数据\"的AI助手，回答基于equipment_health_score.csv等CSV的实时查询，"
  "而非模型训练语料中的过时信息。以下详解每个交互组件。")

doc.add_heading('5.1 SSE流式对话系统', level=2)
P("通信架构采用POST /api/chat → Server-Sent Events流式响应。单次请求可触发多次工具调用（Tool Calling循环），"
  "每次工具调用的结果自动注入对话上下文供模型继续推理。事件类型及前端处理：")
T(["事件类型", "内容", "前端渲染方式", "作用"],
  [["text_delta", "大模型生成的增量文本(每次几字到几十字)", "RAF节流的marked.js Markdown→HTML渲染，追加至当前助手消息气泡", "用户感知的\"打字机\"流式输出"],
   ["tool_call", "模型决定调用的工具名+参数JSON", "可折叠琥珀色卡片(标准工具)或紫色卡片(RAG工具)，标题=工具中文名，内容=参数", "让用户知道AI正在\"查询什么\""],
   ["tool_result", "工具执行返回的结果JSON/文本", "展开tool_call卡片的结果区域，格式化显示", "让用户验证AI查询的结果是否准确"],
   ["chart", "ECharts option JSON配置", "动态创建ECharts实例在助手消息中内联渲染，320px高", "对话中直接看图，无需跳转页面"],
   ["report", "生成的报告文件名+下载链接", "报告下载卡片，含\"查看报告\"和\"下载PDF\"按钮", "对话中触发的报告可一键下载"],
   ["rag_citations", "RAG检索引用列表", "答案末尾紫色引用块，标注来源和匹配度，紫色高亮原文中的引用语句", "每条AI陈述标注知识来源——什么文档、哪个段落"],
   ["done", "对话结束信号", "停止加载动画，恢复输入框，存储对话历史至localStorage", "标记一轮对话完成"]])

doc.add_heading('5.2 推荐问题侧边栏', level=2)
P("17个预编写问题按角色分组展示。问题基于平台真实数据场景设计，确保用户点击即获得有意义的回答。"
  "例如:\"当前哪些设备的Z-Score复合值超过2.0？\"→ AI调用query_z_scores工具→查询z_scores.csv→"
  "返回实际超阈值设备列表。每个问题的回答都调用工具获取实时数据，杜绝硬编码回复。")

doc.add_heading('5.3 RAG三层知识库与引用系统', level=2)
P("AI回答中被RAG知识库支撑的内容标记紫色，点击可查看来源文档全文预览。三级降级确保在任何条件下有检索结果。")
T(["层级", "技术", "延迟", "精度", "适用场景"],
  [["L1 主层", "ChromaDB向量库+BGE中文嵌入\n(bge-large-zh-v1.5, 1024维)", "~15ms", "最高: 语义级理解，可匹配同义改写", "正常运行，回答需要精确引用平台文档"],
   ["L2 降级", "DeepSeek API远程嵌入", "~300ms", "高: 远程大模型语义理解", "BGE本地模型不可用时（如torch未安装）"],
   ["L3 兜底", "TF-IDF关键词匹配", "~5ms", "中: 仅关键词级匹配，无语义泛化", "所有外部服务不可用时的最终保障"]])

doc.add_heading('5.4 Agent决策日志面板', level=2)
P("浮动紫色按钮（右下角）→点击后从底部滑入Timeline面板。展示本轮对话中AI的所有\"思考步骤\"："
  "哪个工具被调用→返回什么结果→AI基于结果做了什么推理→生成了什么回答。"
  "每条日志标注状态（绿色圆点=成功，红色圆点=失败）、调用延迟（ms）和推理摘要。"
  "设计意图：将AI的\"黑箱\"决策过程透明化——让评委和开发者理解AI为什么做出了某个回答。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 六、技术架构总览 ═══════════════════════════════════
doc.add_heading('六、技术架构总览（technical-overview.html）', level=1)
P("技术架构页是平台的技术白皮书，面向开发者角色。左侧浮动TOC（11个锚点），覆盖从系统架构到技术栈的完整技术说明。")

CHART_BLOCK("chart-overview", "系统架构全景图", "ECharts自定义关系图（780px高）",
    "自上而下的4层架构流:\nLayer 1: 角色门(Role Gate)+三个角色图标\nLayer 2: 数据层(4个CSV文件: LOG/SUMMARY/ASSEMBLY/TESTS)\nLayer 3: 五层流水线(Data Prep→Stat→ML→Diagnosis→Decision)\nLayer 4: 应用层(3个角色视图+Dashboard+Chat+Work Orders)",
    "各层之间的数据流箭头连接，展示从原始数据到用户界面的完整路径",
    "静态架构数据（不依赖外部数据文件）", "让开发者一眼理解平台的完整技术架构和模块间依赖关系。",
    "箭头方向=数据流动方向。实线=同步调用，虚线=异步/定时触发。注意Layer 3(流水线)和Layer 4(应用层)之间的双向箭头——"
    "策略切换时前端POST触发后端重算，后端产出的CSV驱动前端渲染",
    "前端ECharts graph类型渲染，节点坐标和连线根据架构逻辑手工布局")

doc.add_heading('6.1 六大算法公式卡片', level=2)
P("6张玻璃态交互卡片，每张包含数学公式（LaTeX风格CSS渲染）、物理直觉文本、实际计算数值。"
  "公式详情参见第一章1.3节\"核心算法公式\"。卡片设计包含hover抬升动画和光晕边框。")

doc.add_heading('6.2 KDE数据天花板交互式演示', level=2)
P("此模块是平台最具教学价值的交互Demo，直接论证\"为什么4参数条件下纯ML不可行\"。")

CHART_BLOCK("kde-v", "电压 KDE密度估计", "ECharts双线KDE图",
    "X轴: 电压值（物理单位V，归一化后）\nY轴: 概率密度",
    "2条曲线:\n- 绿色填充曲线: 正常运行记录的概率密度分布\n- 红色填充曲线: 故障记录的概率密度分布\n两曲线重叠区域=信息论意义上的\"不可区分区域\"",
    "data/kde_params.json — compute_kde_params.py在真实2999行数据上计算: 对正常和故障的电压值分别进行KDE(高斯核, 200点采样)",
    "直观展示电压值在正常和故障状态下的分布重叠程度。重叠越多→仅靠电压值无法区分正常和故障。",
    "两条曲线的重叠面积≈0.95(95%重叠)→Youden's J≈0.077\n"
    "绿色和红色曲线几乎完全重合→\"仅凭电压判断设备是否即将故障\"是不可能的\n"
    "这就是\"KDE天花板\"的含义——信息的物理上限，再好的ML模型也无法突破这个上限",
    "compute_kde_params.py: 从log.csv读取→按Failure.Type=0分正常,≠0分故障→"
    "scipy.stats.gaussian_kde分别拟合→在200个均匀采样点评估密度→导出JSON")

CHART_BLOCK("kde-a", "电流 KDE密度估计", "同上，电流参数", "", "", "data/kde_params.json", "", "同电压KDE，重叠面积≈0.93→Youden's J≈0.064")
CHART_BLOCK("kde-t", "温度 KDE密度估计", "同上，温度参数", "", "", "data/kde_params.json", "", "同电压KDE，重叠面积≈0.92→Youden's J≈0.065")
CHART_BLOCK("kde-r", "转速 KDE密度估计", "同上，转速参数", "", "", "data/kde_params.json", "", "同电压KDE，重叠面积≈0.94→Youden's J≈0.061")

P("右侧Youden's J数字动画：从Phase 0的0.075→Phase 1的0.35→Phase 2的0.65→Phase 3的0.90。"
  "数字颜色从红色(0.075)渐变→琥珀(0.35)→绿色(0.90)。下方滑块控制观察的升级阶段，"
  "每切换阶段时KDE图增至对应的参数维度数量。进度条显示距离\"可用检测水平\"(J=0.60)的进度。"
  "结论卡片随阶段变化更新文字内容。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 七～十五 ═══════════════════════════════════════
# (Sections 7-15 are slightly condensed to keep document manageable)

doc.add_heading('七、工单全流程跟踪看板（work-order-tracking.html）', level=1)
P("工单跟踪页实现维护工单从生成到归档的完整生命周期管理。数据来源: GET /api/work-order-tracking/list。")

doc.add_heading('7.1 6列Kanban看板', level=2)
T(["列", "状态", "卡片颜色标记", "可执行操作"],
  [["待分配\n(Pending)", "工单已生成但未指定技师", "灰色左边框\n无技师信息", "点击→详情面板→\"分配技师\"按钮\n→弹出技师选择弹窗→多选→确认"],
   ["已分配\n(Assigned)", "技师已确认但尚未开始", "蓝色左边框\n显示技师名+电话", "点击→详情面板→\"开始维修\"按钮→状态→执行中"],
   ["执行中\n(In Progress)", "技师正在维修", "琥珀色左边框\n显示开始时间+预计完成", "点击→详情面板→\"提交验收\"按钮→状态→待验收"],
   ["待验收\n(PendingAccept)", "维修完成，等待验收确认", "绿色左边框\n\"通过\"和\"驳回\"按钮", "点击→详情面板→验收通过/驳回\n通过→已完成; 驳回→执行中(重新维修)"],
   ["已完成\n(Completed)", "验收通过，工单关闭", "深绿色左边框\n显示完成时间+验收人", "点击→详情→\"归档\"按钮"],
   ["已归档\n(Archived)", "历史工单，只读保存", "灰色半透明\n不可操作", "仅可查阅"]])

doc.add_heading('7.2 工单卡片内容', level=2)
P("每张Kanban卡片包含以下信息：设备ID（等宽字体加粗）、优先级徽章（P1红脉冲动画 / P2琥珀 / P3蓝）、"
  "维护动作（中文标签）、故障模式（中文标签）、计划停机窗口（时间范围）、健康评分（色码数字）、"
  "成本风险（$金额）、策略标签（彩色边框）、升级徽章（超时工单，红色脉冲）。"
  "卡片拖拽为前端UI交互（预览效果），实际状态变更需通过详情面板的操作按钮和API确认。")

doc.add_heading('7.3 升级机制', level=2)
P("工单超过SLA时限未完成 → notification_service.py自动升级至P1 → 邮件通知主管 → "
  "Kanban卡片出现红色脉冲动画 + \"已升级\"徽章。同一设备72小时内重复生成工单 → 标记\"需根因审查\"。"
  "技师拒绝工单 → 自动重新分配至同技能等级的下一位可用技师。")

doc.add_heading('7.4 技师选择弹窗', level=2)
P("弹出网格布局的技师选择面板。每位技师一张卡片：彩色头像（绿=可用/琥珀=忙碌/灰=离线）、"
  "姓名、类型（初级/高级/电气/热控）、联系方式、负载进度条（绿<60%/琥珀60-85%/红>85%）、"
  "已分配设备标签。支持多选，显示已选数量。确认分配调用POST /api/technicians/assign。")

PAGE_BREAK()

# ═══════════════════════════════════ 八、工作流调度 ═══════════════════════════════════
doc.add_heading('八、工作流调度管理（workflows.html）', level=1)
P("工作流页面管理平台后台的三个自动化定时任务（基于Python APScheduler）。这些任务构成了平台的自主运维中枢——"
  "无需人工干预即可完成每日设备健康巡检、超时工单预警升级和周期性报告分发。以下逐一详解每个组件。")

doc.add_heading('8.1 任务状态卡片（3张）', level=2)
P("三张玻璃态卡片，每张对应一个后台定时任务。数据来源：GET /api/workflows/status，"
  "后端tracking_routes.py从APScheduler的作业存储中读取各任务的最新执行记录。")
T(["任务ID", "中文名称", "调度频率", "功能说明", "状态指示"],
  [["wo_timeout_check", "工单超时检测", "每15分钟", "遍历所有状态为'assigned'或'in_progress'且超过SLA时限的工单 → "
    "自动将状态升级为'escalated' → 调用notification_service.py发送邮件通知主管", "ok(绿)=最近一次检测正常 / fail(红)=检测执行异常 / never(灰)=从未执行"],
   ["daily_health_check", "每日健康巡检", "每天06:00", "从流水线输出目录读取最新的健康评分数据 → 更新equipment_health_score.csv → "
    "统计关键设备数量 → 生成每日健康简报（HTML格式）→ 可选邮件投递", "同上的四态指示"],
   ["weekly_report", "周报生成与分发", "每周一07:00", "汇总本周设备健康趋势、故障统计、维护完成率、成本分析 → "
    "渲染email_weekly_report.html模板 → 生成PDF → 通过SMTP发送给配置的收件人列表 → 存档至reports/generated/", "同上的四态指示"]],
  col_widths=[3.0, 2.5, 2.0, 6.5, 3.5])

P("每张任务卡片显示：任务名称（中文大标题）→ 调度频率描述（副标题）→ 最后运行状态徽章（ok/fail/running/never四态）→ "
  "最后运行耗时（X.X秒）→ 最后执行时间（datetime）→ 结果摘要（截断至40字符）→ 手动触发按钮。"
  "点击手动触发按钮 → POST /api/workflows/trigger/<jobId> → 后端在独立线程中执行任务 → 3秒后自动刷新卡片状态。")

doc.add_heading('8.2 配置展示区', level=2)
P("以标签-数值对的网格展示当前系统运行参数。数据来源：GET /api/workflows/config。"
  "显示字段：当前策略（中文：成本效率/生产效率/质量优先）、每日巡检时间（06:00）、周报生成时间（周一07:00）、"
  "高危设备阈值（health_score<40触发告警，可调参数）、超时升级阈值（24小时，超过此时间未处理的工单自动升级）、"
  "SMTP邮件状态（绿色\"已配置\"或红色\"未配置\"——取决于.env中SMTP_HOST/USER/PASSWORD是否全部设置）。")

doc.add_heading('8.3 执行历史表', level=2)
P("可筛选的日志表格，展示所有定时任务的历史执行记录。数据来源：GET /api/workflows/history?limit=50&job_id=<filter>。")
T(["列名", "内容", "筛选/排序"],
  [["时间", "任务执行开始时间（datetime截断至秒，格式YYYY-MM-DD HH:MM）", "—"],
   ["任务", "任务中文名称（工单超时检测/每日健康巡检/周报生成）", "下拉筛选器：全部任务/单个任务"],
   ["状态", "ok(绿色成功) / fail(红色失败) / running(琥珀色运行中)", "—"],
   ["耗时", "任务执行时长（X.X秒格式）", "—"],
   ["结果", "result_summary或error_message，截断至60字符", "—"]])
P("筛选器：顶部下拉菜单选择\"全部任务\"或单个任务类型 → 自动重新请求API。刷新按钮手动重新加载。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 九、备件库存 ═══════════════════════════════════
doc.add_heading('九、备件库存管理（inventory.html）', level=1)
P("库存管理页面向生产管理负责人，基于(s,S)库存策略模型管理17种CNC备件的采购、存储和消耗。"
  "使用Chart.js 4.4.0（轻量图表库）渲染两个核心图表。页面的数据管道为双重加载："
  "inventory_policy_optimized.csv提供模型参数（再订购点s、目标库存S、EOQ），"
  "GET /api/inventory/stock提供仓库实际库存，后者覆盖前者的估算库存列。"
  "缺货风险在前端重新计算（recomputeRisk函数）：实际库存≤0 → critical(概率0.95)；实际库存<s → high(0.75)；"
  "实际库存<S → medium(0.30)；否则 → low(0.05)。年度缺货成本估算 = (2.5×单价 + 200×交期天数) × 365 × 缺货概率。")

doc.add_heading('9.1 统计卡片（4张）', level=2)
STAT_CARD("inv-stats", "库存统计卡片", "备件种类 | 库存不足数 | 总缺口(件) | 库存总值($)",
    "备件种类=库存数据中不同part_name的数量\n库存不足数=actual_stock < reorder_point_s 的备件数量（琥珀色>0）\n"
    "总缺口=Σ max(0, target_stock_S - actual_stock)\n库存总值=Σ (actual_stock × unit_cost) 前端计算",
    "GET /api/inventory/stock → inventory_connector.check_stock()")

doc.add_heading('9.2 (s,S)库存策略可视化', level=2)

CHART_BLOCK("chart-inv-gap", "当前库存 vs 目标库存缺口对比图", "Chart.js混合柱状+折线图",
    "X轴: 17种备件的中文名称（标签最多60度旋转避免重叠）\nY轴: 数量（件）",
    "数据集1「当前库存(实际)」: 柱状图。每根柱=该备件的当前仓库实际库存。颜色逻辑：实际库存≥目标库存S时绿色柱，<S时琥珀色柱。\n"
    "数据集2「目标库存(S)」: 紫色虚线折线。S由EOQ模型计算——当库存达到S时停止采购。\n"
    "数据集3「重订货点(s)」: 琥珀色虚线折线。当库存降至s时触发采购订单，采购量=S−当前库存。\n"
    "数据集4「缺口(需补货)」: 红色半透明柱状。gap = max(0, S − actual_stock)，仅在当前库存低于S时出现。",
    "模型参数: data/inventory_policy_optimized.csv（s, S, EOQ, unit_cost, lead_time_days, suggested_order_qty）\n"
    "实际库存: GET /api/inventory/stock（覆盖模型中的current_stock_est列）",
    "一眼识别哪些备件需要立即采购补货——红色缺口柱越高的备件越紧迫。紫色虚线(S)和琥珀虚线(s)分别标注了补货的目标上限和安全下限。",
    "红色缺口柱>0 → 该备件当前库存低于目标S → 需要采购补货至S水平。琥珀色柱（当前库存低于s但>0）→ 库存已进入重订货区，虽未用完但趋势下行需关注。"
    "如果柱体超过紫色S虚线 → 库存过剩，可能是采购过多或需求低于预期。比较红色缺口柱的高度可以直观判断各备件的紧急程度排序。",
    "前端: fetch CSV + fetch API → stockMap合并实际库存 → Chart.js注册Bar+Line控制器 → "
    "4个dataset对象（2个bar + 2个line with tension 0.3） → Chart构造函数渲染。每柱颜色通过backgroundColor回调动态判断。")

CHART_BLOCK("chart-risk-donut", "缺货风险分布环形图", "Chart.js环形图（内径cutout: 60%，外 radius: 90%）",
    "环形扇区", "4个扇区（按实际库存计算的缺货风险等级）：\n- Critical（红 #ff453a）: actual_stock ≤ 0\n"
    "- High（琥珀 #ffb340）: 0 < actual_stock < s\n- Medium（蓝 #6db5f9）: s ≤ actual_stock < S\n- Low（绿 #30d158）: actual_stock ≥ S",
    "inventory_policy_optimized.csv → 每备件行 → recomputeRisk()重新计算stockout_risk和probability → 按risk等级分组计数",
    "展示全库17种备件中面临不同程度缺货风险的分布。如果Critical+High占比过高，说明库存策略参数需要整体调整（如提高s/S值或缩短采购周期）。",
    "环形图中心留空——视觉焦点在扇区比例。悬停扇区显示tooltip: \"{label}: {count} 件 ({pct}%)\"。"
    "Critical扇区占比>20%→库存策略严重不足。High+Critical>50%→需整体调高安全库存。如果Low扇区占比>80%→库存充足，可降库存成本。")

doc.add_heading('9.3 缺货风险Top 5表', level=2)
P("HTML表格（非图表库），按risk severity降序排列，取风险最高的5个备件。列：备件中文名（通过PART_CN映射翻译）、"
  "风险徽章（Critical红/High琥珀/Medium蓝/Low绿）、缺货概率（百分比，如\"87%\"）、年缺货成本（$金额，>5000标红）、"
  "建议采购量（件）。排序规则：Critical优先 → 同级按suggested_order_qty降序。"
  "数据来源：同环形图——inventory_policy_optimized.csv经过风险重计算后的前5行。")

doc.add_heading('9.4 采购建议卡片网格', level=2)
P("网格布局（2列CSS grid），每卡片=一个需要采购的备件。筛选条件：actual_stock ≤ reorder_point_s 或 suggested_order_qty > 0。"
  "排序：风险严重度优先 → 同级按缺口大小降序。每卡片显示：紧急程度徽章（\"紧急补货\"红底 / \"建议采购\"蓝底）、"
  "备件中文名+型号、当前库存→缺口→目标库存的对比条（如\"3件 → 缺12件 → 目标15件\"）、EOQ建议订购量、"
  "重订货点s、单位成本和总成本（建议量×单价）。卡片内含库存填充条（fill bar）："
  "宽度=current/S×100%，红色(<30%)/琥珀(30-60%)/绿色(>60%)。紧急卡片左侧有琥珀色粗边框。"
  "汇总栏（网格底部）：涉及备件种类数+总采购件数+总预估成本（Σ(suggested_qty × unit_cost)）。")

doc.add_heading('9.5 库存表格', level=2)
P("完整17行备件清单表，列：零件（中文名+英文原名）、型号（part_number）、库存（点击数字触发prompt()弹出编辑框，"
  "输入新值后POST /api/inventory/adjust）、安全库存（safety_stock=s）、需求（demand_rate）、"
  "缺口（shortage，红色>0）、单价（$）、状态徽章（ok绿/low琥珀/out红）、供应商、操作按钮行（+1/-1/+5/-5快速调整，"
  "每个按钮调用POST /api/inventory/adjust with {part_name, delta}）。数据来源：GET /api/inventory/stock。"
  "刷新按钮重新加载全部数据。\"重新初始化\"按钮POST /api/inventory/stock重置库存快照。")

doc.add_heading('9.6 采购订单管理', level=2)
P("订单卡片列表+自动生成按钮。数据来源：GET /api/inventory/procurement。\"自动生成采购单\"按钮触发"
  "POST /api/inventory/procurement/generate → inventory_connector.generate_procurement_orders()——"
  "基于缺口分析自动生成采购订单。每订单卡片显示：状态徽章（5态：申请中/已下单/运输中/已到货/已入库）、"
  "备件名、数量×供应商、下单日期+预计到货、状态下拉选择器（调用POST /api/inventory/procurement/update-status）、"
  "确认入库按钮（设为\"已入库\"并自动更新库存+记录日志）。")

doc.add_heading('9.7 库存变更日志', level=2)
P("时间戳日志表格。数据来源：GET /api/inventory/logs。列：时间（created_at截断至秒）、零件（中文名）、"
  "变动量（绿色正数表示入库/+N，红色负数表示出库/-N）、原因（reason列）、调整后库存（new_stock）。"
  "每次库存调整（无论是手动编辑还是快速按钮）都会创建一条日志记录。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 十、技师管理 ═══════════════════════════════════
doc.add_heading('十、技师与员工管理（technicians.html）', level=1)
P("技师管理页面面向生产管理负责人，管理维修团队的完整信息。支持5种技师类型：初级技师(junior_technician)、"
  "高级技师(senior_technician)、电气专员(electrical_specialist)、热控专员(thermal_specialist)、"
  "机械专员(mechanical_specialist)。三种工作状态：在岗(available)、忙碌(busy)、休假(off_duty)。")

doc.add_heading('10.1 统计卡片（4张）', level=2)
STAT_CARD("tech-stats", "技师统计卡片", "员工总数 | 在岗人数 | 忙碌人数 | 当前工单总数",
    "员工总数=technicians数组length\n在岗人数=status='available'计数（绿色）\n忙碌人数=status='busy'计数（琥珀色）\n"
    "当前工单总数=SUM(current_workload) 所有技师的当前分配工单之和",
    "GET /api/technicians/workloads/summary → tracking_routes.py")

doc.add_heading('10.2 技师名册表', level=2)
P("可按类型和状态双重筛选的HTML表格。数据来源：GET /api/technicians。筛选器：类型下拉（全部/电气/热控/高级/初级/机械）"
  "+ 状态下拉（全部/在岗/忙碌/休假）。")
T(["列名", "内容", "交互"],
  [["ID", "技师编号（如T001）", "—"],
   ["姓名", "中文姓名（加粗显示）", "—"],
   ["类型", "中文类型标签（TYPE_CN映射）", "按类型筛选"],
   ["邮箱/电话", "联系方式（小号字体）", "—"],
   ["状态", "徽章：在岗(绿)/忙碌(琥珀)/休假(灰)", "按状态筛选"],
   ["工单负载", "\"当前数 / 最大负载\" + 颜色编码\n红=≥100%, 琥珀=70-99%, 绿=<70%", "—"],
   ["操作", "编辑(打开预填弹窗) / 切换状态(PUT) / 删除(DELETE)", "删除仅限无活跃工单的技师"]])
P("新增员工按钮打开空白表单弹窗。")

doc.add_heading('10.3 负载仪表盘', level=2)
P("网格布局的技师卡片（最小宽310px，自适应列数）。每张卡片：彩色头像圆圈（姓名首字，绿/琥珀/灰对应状态）、"
  "姓名+类型+状态标签、联系方式（邮箱 | 电话）、负载进度条（横向填充条，宽度=负载/最大×100%，"
  "绿色<70% / 琥珀70-99% / 红色≥100%）、已分配设备标签列表（或\"暂无工单\"）、编辑/切换/删除操作按钮。"
  "数据来源：GET /api/technicians（无筛选，全量返回）。")

doc.add_heading('10.4 新增/编辑弹窗', level=2)
P("7字段表单：姓名（必填，placeholder\"张伟\"）、技师类型（5选1下拉）、邮箱、电话、最大负载（数字1-10，默认3）、"
  "技能标签（文本，逗号分隔，如\"高电压,电路诊断\"）、状态（在岗/忙碌/休假）。"
  "保存时：新增→POST /api/technicians；编辑→PUT /api/technicians/<id>。保存成功后自动刷新名册和负载仪表盘。")

doc.add_heading('10.5 技能匹配规则', level=2)
P("维修工单的技师自动分配由technician_assigner.py实现，遵循优先级规则链：")
T(["设备状态", "异常模式", "分配技师类型", "优先级理由"],
  [["ALARM", "voltage_drift", "electrical_specialist（电气专员）", "电压漂移涉及电路/电源系统，需要电气专业诊断"],
   ["ALARM", "thermal_buildup", "thermal_specialist（热控专员）", "热积聚涉及冷却/散热系统，需要热控专业诊断"],
   ["ALARM", "combined_degradation", "senior_technician（高级技师）", "多参数同步恶化需要综合判断，高级技师经验更丰富"],
   ["ALARM", "power_anomaly", "electrical_specialist", "功率异常通常由电路问题导致"],
   ["WARNING/WATCH", "voltage_drift/thermal_buildup", "junior_technician（初级技师）", "预警级别较低，初级技师可以处理"],
   ["任意", "normal", "junior_technician", "例行检查，无需专业技师"]],
  col_widths=[2.0, 3.0, 4.0, 7.0])
P("分配时附加负载均衡：同类型技师中优先选择当前负载率最低的。如果所有同类型技师均已满载（负载率≥100%），自动升级为高级技师。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 十一、知识库 ═══════════════════════════════════
doc.add_heading('十一、知识库管理中心（knowledge-base.html）', level=1)
P("知识库管理页面向平台开发人员，提供ChromaDB向量数据库的完整管理界面。页面采用双栏布局（左侧240px侧边栏 + 主内容区），"
  "集成文档上传、语义检索测试、RAG效果对比、三级降级验证和检索日志五大功能模块。"
  "底层技术栈：ChromaDB（向量存储）+ BGE中文嵌入模型（bge-large-zh-v1.5）+ DeepSeek API（远程嵌入备用）+ TF-IDF（兜底）。")

doc.add_heading('11.1 左侧栏 — 文档集合管理器', level=2)
P("三个玻璃态集合卡片，点击切换当前操作的文档集合：")
T(["集合名",  "集合ID", "内容说明", "文档数/分块数", "上传功能"],
  [["系统文档库", "sys_docs", "平台技术文档、算法说明、API文档、架构设计文档的向量化存储", "~8文档/~80+分块", "无（自动生成）"],
   ["运维知识库", "maint_kb", "故障模式库、维修操作手册、巡检标准、安全规范——运维人员可上传补充", "~12文档/~120+分块", "支持拖拽/点击上传(.md/.txt/.pdf/.docx, ≤10MB)"],
   ["故障案例库", "fault_cases", "历史维修记录、根因分析、处置效果评估——自动从流水线输出生成", "~6文档/~60+分块", "无（自动生成）"]],
  col_widths=[2.5, 2.0, 5.5, 3.0, 3.5])
P("激活集合卡片 → 紫色边框+紫色背景发光。状态指示灯（集合卡片下方）：绿色圆点=ChromaDB引擎就绪，红色圆点=引擎不可用。"
  "引擎状态由GET /api/knowledge-base/stats的engine_available字段决定。")

doc.add_heading('11.2 主面板 — 文档列表', level=2)
P("标题\"文档列表（N个，共M块）\"。操作按钮行：上传文档（仅maint_kb集合可见，点击触发文件选择器，"
  "支持拖拽到上传区域）、重新索引当前集合（POST /api/knowledge-base/reindex with {collection}——"
  "清空ChromaDB中的该集合并重新嵌入所有文档）、重建全部索引（POST /api/knowledge-base/rebuild-all——"
  "清空全部三个集合并从头嵌入，耗时较长约30-60秒）。")
P("文档表格列：文档名称（等宽字体，hove显示完整路径）、大小（KB）、状态（\"已索引\"绿色标签 / \"未索引\"琥珀色标签）、"
  "操作（\"自动生成\"灰标签 或 \"重新索引\"+\"删除\"按钮）。删除操作调用DELETE /api/knowledge-base/documents with {source, collection}——"
  "从ChromaDB中删除该文档的所有chunk，同时从文件系统中删除源文件。")

doc.add_heading('11.3 主面板 — 语义检索测试', level=2)
P("搜索框 + 集合选择器（全部知识库/sys_docs/maint_kb/fault_cases）+ \"搜索\"按钮。"
  "调用GET /api/knowledge-base/search?q=<query>&k=5&collection=<coll> → kb_routes.py → rag_engine.search()（指定集合）"
  "或search_all()（全部集合）。每条搜索结果显示：[集合名] 文档名 > 段落标签 + 匹配度得分（3位小数）+ "
  "内容摘要（截断300字符）。结果可点击 → 弹出引用预览浮层，显示完整文档内容（含格式保留的pre标签）。"
  "摘要行显示\"共N条结果，耗时Xms\"。")

doc.add_heading('11.4 主面板 — RAG效果对比演示', level=2)
P("此模块是知识库页面的核心亮点——直观展示RAG（检索增强生成）对AI回答质量的提升效果。"
  "三个预设问题按钮：\"健康分怎么计算？\"（默认选中）、\"Type 4故障怎么处理？\"、\"如何切换维护策略？\"。"
  "左右双栏对比布局：")
T(["", "左侧：无RAG", "右侧：有RAG（本项目RAG系统）"],
  [["边框颜色", "红色边框（警告色）", "绿色边框（成功色）"],
   ["标签", "\"无 RAG（纯大模型通用回答）\"", "\"有 RAG（本项目 RAG 系统）\""],
   ["回答特征", "通用、模糊、无具体数字\n如\"健康评分通常基于多维度指标...\"", "具体、可追溯、含真实数字\n如\"健康评分H=100×exp(-k×R)，其中R为融合风险评分，当前CNC_036健康分=32\""],
   ["引用标注", "无", "紫色引用徽章[1][2]...，标注来源文档+段落+匹配度\n点击可查看原文预览"],
   ["数据来源", "DeepSeek模型通用知识", "本项目CLAUDE.md、算法文档、CSV数据统计\n通过ChromaDB+BGE语义检索注入"]],
  col_widths=[2.0, 7.0, 7.5])
P("每个预设问题的对比内容（noRag和withRag文字）预硬编码在JS的RAG_PRESETS数组中，确保评委在看Demo时零延迟加载。"
  "引用预览浮层（citePreviewOverlay）不仅服务于预设问题，也服务于实际搜索结果——显示来源标题、路径、匹配度得分、"
  "完整文档内容（最大高度60vh可滚动）。Esc键、X按钮或点击浮层外部均可关闭。")

doc.add_heading('11.5 降级测试面板', level=2)
P("此面板用于验证三级RAG降级链的实际效果。3个单选按钮选择嵌入层级：BGE本地（~15ms，默认）/ DeepSeek API（~300ms）/ "
  "TF-IDF（~5ms）。查询输入框+\"测试\"按钮（POST /api/knowledge-base/degrade-test with {query} → "
  "返回指定层级的检索结果）或\"三级对比测试\"按钮（同时运行三个层级并排展示结果）。"
  "结果三栏布局（仅在三级对比模式）：BGE本地(绿色标题) | DeepSeek API(琥珀标题) | TF-IDF(红色标题)。"
  "每栏显示检索结果的段落标题、匹配度徽章（≥70%绿色高匹配 / 40-69%琥珀中等 / <40%红色低匹配）、"
  "来源文档名、内容片段（120字符）。目的是可视化三级降级的检索质量梯度——BGE质量最高，TF-IDF虽有信息损失但仍可返回关键词相关结果。")

doc.add_heading('11.6 检索日志', level=2)
P("实时展示最近50条RAG检索调用记录。数据来源：GET /api/knowledge-base/logs?limit=50。"
  "每条日志：时间戳 → 查询文本（截断）→ 目标集合 → 结果数（N条）→ 耗时（Xms）。自动刷新。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 十二、报告中心 ═══════════════════════════════════
doc.add_heading('十二、报告中心（reports.html）', level=1)
P("报告中心提供历史报告的浏览、搜索、类型筛选、HTML预览和PDF转换功能。报告文件存储在web-dashboard/reports/generated/目录"
  "（HTML格式）和reports/pdfs/目录（PDF格式），由定时任务（weekly_report）或手动触发生成。"
  "报告模板采用Apple风格UI设计（iOS26玻璃卡片+环境光晕+打印兼容CSS）。")

doc.add_heading('12.1 报告类型识别', level=2)
P("后端app.py根据中文文件名前缀自动判定报告类型（5种）：")
T(["报告类型", "类型标识", "中文前缀匹配", "主要内容"],
  [["周度系统报告", "weekly", "\"周度系统报告\"", "设备健康全景、故障趋势、维护完成率、本周成本统计、下周预测"],
   ["单设备报告", "device", "\"单设备报告\"", "单台设备完整健康档案、历史趋势折线图、SHAP归因分析、维护建议"],
   ["风险告警报告", "risk", "\"高风险设备报告\"", "关键设备健康详情、根因分析、多信号评估、建议紧急措施、升级状态"],
   ["温度专项报告", "thermal", "\"热漂移分析报告\"", "热积聚分析、温度时序趋势、冷却系统评估、红外检测建议"],
   ["低健康分报告", "health_critical", "\"低健康分报告\"", "健康评分<40的设备汇总、逐设备退化分析、批量维护建议"],
   ["通用报告", "general", "（无匹配前缀的fallback）", "自定义数据范围，内容可配置"]],
  col_widths=[2.5, 2.5, 3.5, 7.5])

doc.add_heading('12.2 统计卡片行（4张）', level=2)
STAT_CARD("report-stats", "报告统计卡片", "报告总数 | 报告类型数 | 总大小(MB) | 含PDF数",
    "报告总数=reports/generated/目录中.html文件数量\n报告类型数=去重后的report_type值数\n"
    "总大小=Σ 各.html文件大小÷1024（MB）\n含PDF数=有对应.pdf文件的报告数量",
    "GET /api/reports → app.py扫描文件系统")

doc.add_heading('12.3 工具栏', level=2)
P("搜索框（按文件名或报告类型子串过滤，不区分大小写）+ 类型筛选下拉（全部/周报/风险/设备/温度/通用）+ "
  "排序按钮（循环8种排序：时间降序/升序→名称降序/升序→大小降序/升序→类型降序/升序，按钮显示当前排序模式标签）+ "
  "刷新按钮 + \"打开PDF文件夹\"按钮（POST /api/reports/open-pdfs-folder → 在操作系统文件管理器中打开PDF目录）。")

doc.add_heading('12.4 报告卡片列表', level=2)
P("每个报告渲染为一张玻璃态卡片：左侧类型emoji图标（周刊=书本📗绿/风险=警告⚠红/设备=齿轮⚙青/温度=温度计🌡红，"
  "颜色按类型编码）→ 文件名（等宽字体，自动换行）→ 元数据行（类型标签+文件大小KB/MB+格式化时间）→ "
  "操作按钮行：「查看」链接（新标签页打开HTML）| 「PDF」下载（如已生成PDF，绿色按钮）| "
  "「生成PDF」（如尚未生成，调用POST /api/reports/generate-pdf → Playwright无头Chromium渲染HTML→A4 PDF → 按钮变为下载链接）"
  "| 「删除」（红色，POST /api/reports/delete → 删除HTML+对应PDF → 刷新列表）。")

doc.add_heading('12.5 角色过滤', level=2)
P("不同角色看到的报告类型不同（ROLE_REPORT_TYPES配置）：运维工程师看到device/thermal/parts_summary/work_order类报告"
  "（关注具体设备和备件）；生产管理负责人看到weekly/risk/health_critical/sensor_advisory类报告（关注趋势和风险）；"
  "平台开发人员看到全部类型。前端的filterReports函数根据当前角色过滤reportList。")

doc.add_heading('12.6 邮件投递系统', level=2)
P("report_delivery.py负责报告的分发。deliver_report()函数：接收报告规格(ReportSpec)+渲染HTML+报告类型→"
  "写入reports/generated/（HTML）和reports/pdfs/（PDF）→可选Send Email（通过notification_service._send_email()）→"
  "更新spec.export_meta（URL、时间戳、文件大小）→返回DeliveryResult（success、URLs、sizes、errors）。"
  "文件命名规则：周报/风险/温度类→\"{中文类型名}_{YYYYMMDD_HHMMSS}.html\"；设备类→\"{类型名}_{设备ID}_{YYYYMMDD_HHMMSS}.html\"。"
  "邮件通知系统（notification_service.py）基于SMTP配置（环境变量SMTP_HOST/PORT/USER/PASSWORD/FROM），"
  "支持HTML+纯文本双格式（MIME multipart/alternative），暗色标题栏+玻璃卡片风格的邮件正文模板。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 十三、设备矩阵 ═══════════════════════════════════
doc.add_heading('十三、独立设备健康矩阵（device-grid.html）', level=1)
P("独立设备矩阵页（路由 /device-grid）是首页10×10网格的纯净全屏版本。去除首页的KPI面板、管理策略对比、"
  "故障注入演示和Demo HUD等复杂元素，专注于100台设备的健康状态可视化。适合作为日常运维的主监控视图或大屏展示。"
  "与首页共享device-grid-component.js核心组件，网格渲染、详情面板、SHAP探索、基线追溯的行为完全一致。")

doc.add_heading('13.1 与首页grid的关键差异', level=2)
T(["方面", "home.html中的grid", "device-grid.html"],
  [["页面复杂度", "高——含KPI卡片、管理视图、故障注入、Demo模式", "低——纯网格视图，加载更快（373行HTML）"],
   ["顶部统计条", "无（KPI卡片代替）", "有——网格上方统计条显示各健康等级设备数量+活跃工单数"],
   ["图例说明", "无（隐含在配色中）", "有——6项图例：健康绿/警告琥珀/退化橙/关键红/工单脉冲/优先级高亮"],
   ["角色视图", "运维和管理看到不同内容", "所有角色看到相同内容（纯设备监控）"],
   ["交互功能", "完整（详情面板+SHAP+基线+故障注入+Demo）", "完整（详情面板+SHAP+基线，无故障注入和Demo）"],
   ["适用场景", "综合入口页面", "日常监控/大屏展示/运维值班"]],
  col_widths=[3.0, 6.0, 7.0])

doc.add_heading('13.2 图例条', level=2)
P("网格上方6项图例（grid-legend）：绿色圆点=健康(Healthy) | 黄色圆点=警告(Warning) | 橙色圆点=退化(Degrading) | "
  "红色圆点=关键(Critical) | 红色脉冲圆点=有开放工单(ALARM级脉冲动画) | 琥珀脉冲边框=Top 3成本风险设备(金色优先脉冲)。"
  "图例帮助不熟悉颜色编码的评委或新用户快速理解网格含义。")

doc.add_heading('13.3 详情面板内容（与首页共用逻辑）', level=2)
P("点击任意设备单元格打开右侧详情面板（420px宽，玻璃态背景），分为六个信息区：")
T(["面板区域", "显示内容", "数据来源"],
  [["① 设备头部", "设备ID + 健康评分（色码大数字）+ 健康等级标签 + 趋势箭头 + 工单告警级别", "equipment_health_score.csv"],
   ["② 关键指标", "4个迷你指标：健康评分(色码)、ML故障密度(%)、维护超期(天)、成本风险($k)", "equipment_health_score.csv + cost_risk_matrix.csv"],
   ["③ 异常信号", "每个激活的异常信号一行：严重度圆点(红/琥珀/绿)、特征名、当前值、解释文本", "z_scores.csv + diagnosis_results.csv"],
   ["④ SHAP归因", "自然语言摘要 + Top 5贡献特征（特征名+类别标签+SHAP值颜色码+箭头方向）+\"交互式探索\"按钮", "shap_dashboard.json"],
   ["⑤ 维护工单", "优先级(P1-P3)+维护动作(中文)+紧急度评分+停机窗口+预期节约($)+根因(Top2风险因子+SHAP总结)", "maintenance_work_orders.csv"],
   ["⑥ 检查清单", "编号列表项（青色序号）：外观检查4项+电气检查3项+备件替换2项，可逐项勾选", "预定义的CNC维护标准清单"]],
  col_widths=[2.5, 8.0, 5.5])
P("面板底部有\"加入工单跟踪\"按钮，调用POST /api/work-order/create with {machine_id, strategy}。"
  "系统检测重复工单和跨策略转移，弹出确认提示。")

doc.add_heading('13.4 SHAP探索弹窗', level=2)
P("点击详情面板中的\"交互式探索\"按钮或grid的SHAP入口，打开SHAP全局探索弹窗。"
  "特征选择下拉 → ECharts散点图（460px高，X轴=特征原始值带单位，Y轴=SHAP贡献值，"
  "每点=一台设备，颜色=健康等级，大小=风险评分）。tooltip显示设备ID、特征值、SHAP贡献、健康等级。"
  "底部统计条：设备总数、风险升高数(红)、风险降低数(绿)、平均贡献值。标签行显示受显著影响的设备数和最大贡献。"
  "数据来源：data/shap_scatter_data.json。关闭方式：X按钮/Esc键/点击浮层外部。")

doc.add_heading('13.5 基线追溯面板', level=2)
P("右键点击网格中的设备单元格打开三选项卡追溯弹窗：")
T(["选项卡", "内容", "数据来源", "解读要点"],
  [["Health 健康追溯", "8维度健康分加权分解（每个维度：标签+权重%+原始值+贡献分）\n水平条形图可视化每维度贡献（红→琥珀→灰）\n基线来源标注（自基线/混合/集群回退）+基线质量\n自然语言解释（趋势+主要风险因子）", "baseline_stats.csv", "红色条形=该维度是健康评分低的主因\n基线来源为\"集群回退\"→该设备基线不够精确"],
   ["Z-Score 统计追溯", "3参数表格：电压/电流/温度的当前值、基线均值、标准差、Z-Score值\n复合Z-Score+颜色指示\n标注：Z>2.0触发告警，精度83.9%", "z_scores.csv", "Z>2.0的参数字段=异常主因\n复合Z-Score综合判断多参数偏离程度"],
   ["Risk 风险追溯", "成本风险=故障概率×单位成本×日产量\n日成本>$5k标红警告", "cost_risk_matrix.csv", "日成本风险超过$5k需立即排程\n将统计概率映射为管理层可理解的金额"]],
  col_widths=[3.0, 6.5, 3.5, 3.5])

PAGE_BREAK()

# ═══════════════════════════════════════════ 十四、鹰眼球体 ═══════════════════════════════════
doc.add_heading('十四、鹰眼球体 — 3D数字孪生（sphere-demo.html）', level=1)
P("鹰眼球体是平台的3D数字孪生演示页面（路由 /sphere-demo），使用Three.js 0.160.0（ES Module导入）"
  "在WebGL画布上渲染100台CNC设备的三维空间数字孪生。页面不加载导航栏（使用design-tokens.css代替），"
  "全屏Canvas + 叠加UI面板的沉浸式设计。支持鼠标、触摸和手势（MediaPipe HandLandmarker）三种交互方式。")

doc.add_heading('14.1 Three.js场景组件', level=2)
T(["组件", "技术实现", "参数", "视觉效果"],
  [["渲染器", "WebGLRenderer + ACESFilmicToneMapping", "曝光1.2, 像素比≤2, 透明背景", "电影级色调映射，高光不爆白"],
   ["灯光系统", "AmbientLight(蓝调,强度2.5) +\nDirectionalLight key(蓝白,强度4,右上) +\nDirectionalLight fill(橙红,强度1.5,左下)", "三点照明体系", "球体向阳面偏蓝、背阴面偏暖——增强立体感"],
   ["星空背景", "500个随机粒子点(120单位立方体)", "PointsMaterial, 蓝色小点, 透明度0.6", "深空科技感背景"],
   ["设备球体", "100张CanvasTexture卡片按Golden Spiral排列\nφ=acos(1-2(i+0.5)/n), θ=π(1+√5)i, r=15", "卡片1.6×1.0 PlaneGeometry\n256×160px OffscreenCanvas", "完美均匀分布的3D球面布局——Fibonacci球"],
   ["卡片纹理", "每卡片Canvas绘制：圆角裁剪+暗玻璃底(#0d0d14)\n+健康色4px描边+发光顶线\n+设备ID(白色粗体48px)+HP分数(彩色20px)", "180度翻转后贴图", "设备ID和健康评分清晰可读，颜色编码与2D页面一致"],
   ["卡片材质", "MeshStandardMaterial +\nemissive=健康色×8%强度 +\nDoubleSide渲染", "金属粗糙度0.4", "卡片发出微弱的与健康颜色匹配的环境光"],
   ["轨道环1", "Torus(R+0.3, 青色, XZ平面)", "透明度0.25, 慢速自转", "Siri-Orb风格的视觉装饰"],
   ["轨道环2", "Torus(R+0.35, 紫色, 倾斜PI/3 X + PI/4 Y)", "透明度0.15, 反向慢转", "增加视觉层次和科技感"]],
  col_widths=[2.0, 5.0, 3.5, 5.5])

doc.add_heading('14.2 UI叠加面板', level=2)

doc.add_heading('信息面板（左上）', level=3)
P("玻璃态半透明面板。标题\"鹰眼\"（青→紫渐变文字，STXingkai华文行楷字体）+ 副标题\"3D工业数字孪生·手势交互\"。"
  "2×2统计网格：健康(绿)/警告(琥珀)/退化(橙)/危险(红)四个实时计数——数据来自equipment_health_score.csv的health_status列。"
  "选中设备信息区（点击球面卡片后更新）：设备ID、健康评分（色码大数字）、健康等级、主要风险因子。")

doc.add_heading('手势指示器（底部居中）', level=3)
P("半透明药丸状状态栏。显示当前交互模式：\"鼠标模式——拖拽旋转·滚轮缩放\"或\"手势模式\"（手势激活时绿色高亮）。")

doc.add_heading('摄像头预览（右下角）', level=3)
P("160×120px圆角窗口，CSS scaleX(-1)镜像翻转。WebCam实时画面，半透明(opacity 0.6)，hover时完全不透明。"
  "为MediaPipe手势识别提供输入源。如果摄像头初始化失败，弹出居中提示\"摄像头不可用\"（含错误信息和\"知道了\"关闭按钮）。")

doc.add_heading('调试面板（右上角）', level=3)
P("深色半透明覆层，标题\"GESTURE PIPELINE\"（青色大写）。8行诊断信息（每30帧更新一次）："
  "Camera（OK/WAIT+分辨率，绿/琥珀）| Tick loop（帧计数）| send() ok/err（成功/失败比，错误多时红色）| "
  "onResults（手部检测结果接收数）| Hand count（当前检测到的手数）| Gesture（当前手势类型：none/pan/pinch/twist）"
  "| rotY/rotX（当前旋转角度，2位小数）| scale（当前缩放比例，2位小数）。用于开发调试和评委展示手势识别管线。")

doc.add_heading('14.3 手势识别管道', level=2)
P("平台集成MediaPipe HandLandmarker（Tasks Vision 0.10.14）实现无接触手势操控3D球体——"
  "设计意图是评审现场演示者无需接触键盘鼠标即可操控。管道流程：")
T(["阶段", "操作", "技术细节"],
  [["① 摄像头初始化", "getUserMedia(640×480, 前置摄像头) → 动态import @mediapipe/tasks-vision → FilesetResolver.forVisionTasks()加载WASM → 加载hand_landmarker.task模型文件", "模型存储在/data/hand_landmarker.task"],
   ["② 手部检测", "HandLandmarker.detectForVideo(video, timestamp) per frame", "VIDEO模式，最多4只手，检测/存在/追踪置信度均为0.3"],
   ["③ 手势分类", "0手→无手势；1手（锁定3帧）→Pan平移；2手（锁定8帧）→Pinch缩放+Twist旋转", "锁定帧机制防止误触发（短暂经过摄像头的手不会触发操控）"],
   ["④ 数据平滑", "EMA指数移动平均 → 当前值×0.7 + 新值×0.3", "防止手势抖动，交互更流畅"],
   ["⑤ 场景变换", "Pan→球体旋转(灵敏度3.0x)；Pinch→缩放(范围0.4-3.0,每帧±8%)；Twist→Z轴旋转(灵敏度2.0x)", "旋转角X限制±90°防止翻转"]],
  col_widths=[2.5, 5.5, 8.0])

doc.add_heading('14.4 鼠标/触摸交互（回退方案）', level=2)
P("当手势不可用时（摄像头未授权/光线不足/浏览器不兼容），自动回退到传统交互："
  "鼠标拖拽（左键按住旋转，灵敏度0.005）+ 鼠标滚轮（缩放因子0.92/1.08，范围0.4-3.0）+ "
  "点击射线检测（Raycaster与100张卡片求交→选中设备→打开详情面板）。"
  "触摸：单指拖拽旋转。惯性衰减（每帧速率×0.95）。空闲自动旋转（绕Y轴0.0015 rad/帧）。"
  "所有交互使用8.0阻尼因子的平滑插值。")

doc.add_heading('14.5 共享组件集成', level=2)
P("球体场景集成了device-grid-component.js的完整交互：点击卡片→openPanel()打开详情面板（同首页/设备矩阵）、"
  "详情面板中的\"交互式探索\"按钮→openShapExploration()打开SHAP散点图弹窗。"
  "数据加载：PapaParse解析equipment_health_score.csv → 构建100个设备对象 → 为每设备生成Canvas纹理。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 十五、评委讲解助手 ═══════════════════════════════════
doc.add_heading('十五、评委讲解助手（assistant.js）', level=1)
P("评委讲解助手是注入在5个核心页面的共享智能模块（通过<script type=\"module\">加载），以浮动3D球体按钮和话术库系统"
  "为评委/用户提供每个页面每个板块的即时讲解。设计灵感来源于Apple Siri的球体动画。"
  "核心设计理念：话术库优先（零延迟预生成讲解）→ AI回退（未命中时SSE流式生成定制讲解）→ 实时数据注入（模板变量替换）。")

doc.add_heading('15.1 3D球体按钮', level=2)
P("右下角56px青色渐变球体（assistant-orb）：径向渐变模拟3D球体明暗 + conic-gradient轨道环（5秒旋转动画）+ "
  "镜面高光点 + 呼吸动画（3秒ease-in-out，缩放1.0→1.08）→ 悬停时显示三层脉冲扩散环。"
  "球体可拖拽移动（鼠标和触摸均支持），4px移动阈值区分点击vs拖拽。位置持久化到localStorage（assistant-orb-left/top）。"
  "点击打开讲解面板 → 面板跟随球体位置（优先显示在球体上方，空间不足时显示在下方）。")

doc.add_heading('15.2 话术库结构', level=2)
P("SPEECH_LIBRARY对象以页面文件名（如\"home.html\"、\"index.html\"）为主键，每个页面包含多个section条目。每条话术的结构：")
T(["字段", "内容", "用途"],
  [["title", "中文板块标题（如\"设备健康矩阵\"、\"KDE数据天花板\"）", "面板中显示为讲解主题"],
   ["summary", "一句话概述", "快速了解该板块的核心功能"],
   ["speech.what", "\"这个模块是做什么的\"——功能描述段落（150-250字）", "回答评委\"这是什么\"的问题"],
   ["speech.why", "\"为什么要这样设计\"——设计动机和解决的问题", "回答评委\"为什么\"的问题"],
   ["speech.effect", "\"效果如何\"——量化效果和指标", "回答评委\"有什么用\"的问题"],
   ["speech.relation", "\"和项目整体的关系是什么\"——在平台中的定位", "回答评委\"关联什么\"的问题"],
   ["keywords", "中文关键词数组（如[\"健康评分\",\"设备状态\",\"颜色编码\"]）", "用于匹配引擎的关键词命中"]],
  col_widths=[3.0, 5.5, 7.0])
P("话术库覆盖全部13个页面、40+板块。每条话术的定量部分包含模板变量（如${criticalCount}、${meanScore}），"
  "在渲染时从GET /api/health-summary获取实时数据替换——确保话术中的数字始终反映平台当前实际状态。")

doc.add_heading('15.3 匹配引擎', level=2)
P("matchSpeech()函数实现五级优先级匹配：")
T(["优先级", "匹配方式", "触发条件", "响应速度"],
  [["1（最高）", "精确匹配", "页面上某个元素有data-speech属性，属性值正好匹配话术库中的section key", "零延迟"],
   ["2", "URL Hash匹配", "URL中的#锚点（如#sec-kde）对应话术库中的section key", "零延迟"],
   ["3", "标题关键词匹配", "当前视口中可见的标题文字匹配话术条目keywords数组中≥2个关键词", "零延迟"],
   ["4", "用户问题匹配", "用户在弹窗输入框中提问，问题文本匹配话术keywords中≥1个关键词", "零延迟"],
   ["5（回退）", "页面默认概述", "无任何匹配时，返回该页面的__default__条目（各页面通用概述）", "零延迟"],
   ["6（AI回退）", "SSE流式AI生成", "以上全部未命中 → POST /api/assistant/explain → SSE流式 → DeepSeek生成", "~2-5秒"]],
  col_widths=[2.0, 3.0, 6.5, 2.0])

doc.add_heading('15.4 面板状态', level=2)
P("讲解面板（白磨砂玻璃风格，最大宽400px）有五种显示状态：")
T(["状态", "触发条件", "面板内容", "底部按钮"],
  [["匹配成功", "话术库命中", "四段式布局：做什么📋/为什么重要💡/效果如何📊/和项目关系🔗，每段带emoji标题+正文", "复制 | 简化一点 | 展开讲细 | 追问 | 关闭"],
   ["AI流式生成", "话术库未命中+AI端点可用", "流式文字+闪烁光标（SSE text_delta事件）\n**粗体**渲染为<strong>", "复制 | 追问 | 关闭"],
   ["空状态", "话术库未命中+AI未调用", "\"暂未找到匹配的讲解话术\"提示+\"试试输入：解释这个模块\"建议", "手动提问 | 关闭"],
   ["错误状态", "AI端点调用失败", "警告图标+\"AI生成失败\"+跳转AI Copilot链接", "重试 | 关闭"],
   ["追问输入", "用户点击\"追问\"按钮", "动态textarea输入框+\"发送\"/\"取消\"按钮，提交后重新运行匹配引擎", "发送 | 取消"]],
  col_widths=[2.0, 4.0, 5.5, 4.5])

doc.add_heading('15.5 AI回退流程', level=2)
P("当话术库5级匹配全部未命中时，触发AI回退：POST /api/assistant/explain with {context, question, mode} → "
  "后端routes.py构建包含健康上下文和模式指令的system prompt → DeepSeek API via httpx流式生成 → "
  "SSE text_delta事件返回 → 前端逐字追加渲染。可选的mode参数：simplify（简化当前话术）、expand（展开深入讲解）、"
  "默认（生成新讲解）。简化/展开模式同时传递previous_text（当前话术原文，截断1500字符）供AI参考。"
  "模板变量从前端预获取的/api/health-summary数据中替换——如${meanScore}替换为\"67.3\"、${criticalCount}替换为\"12\"等。")

doc.add_heading('15.6 跨页面集成', level=2)
P("助手注入在5个核心页面（home.html、index.html、chat.html、technical-overview.html、device-grid.html）中。"
  "AssistantPanel类在页面加载完成后自动初始化。通过window._assistantQuickAsk(text)方法，"
  "AI Copilot页面可将用户问题（如\"解释一下这个KDE图\"）传递到助手面板中直接匹配或生成讲解。")

PAGE_BREAK()

# ═══════════════════════════════════════════ 附录 ═══════════════════════════════════════
doc.add_heading('附录A：核心数据文件清单', level=1)
T(["文件", "类型", "行/大小", "生产者", "消费者", "核心字段"],
  [["equipment_health_score.csv", "CSV", "100行", "stat_inference", "home/device-grid/sphere-demo/index",
    "machine_id, health_score, health_status, health_level, failure_mode, cost_at_risk, risk_tier"],
   ["maintenance_work_orders.csv", "CSV", "~30-60行(策略相关)", "decision", "home/index/work-order-tracking",
    "machine_id, priority(P1/P2/P3), action, fault_pattern, technician, spare_parts, downtime_window, estimated_cost, expected_savings, sla_hours"],
   ["shap_dashboard.json", "JSON", "100对象", "ml_inference(SHAP)", "home/device-grid",
    "per device: {shap_values: {Voltage, Amperage, Temperature, RotorSpeed}, top_feature, mean_shap, category_breakdown}"],
   ["z_scores.csv", "CSV", "~3000行(100设备×30天)", "data_prep", "index/home/device-grid-component/routes.py",
    "Date, Equipment.Id, z_Voltage, z_Amperage, z_Temperature, z_composite"],
   ["baseline_stats.csv", "CSV", "100行", "stat_inference", "device-grid-component/index",
    "Equipment.Id, V_mean, V_std, A_mean, A_std, T_mean, T_std, baseline_source(Self/Mixed/Cluster)"],
   ["algorithm_comparison.csv", "CSV", "7行(7算法)", "benchmark_algorithms.py", "index/technical-overview",
    "algorithm, type, AUC, Accuracy, Precision, Recall, F1, training_time_ms"],
   ["kde_params.json", "JSON", "4参数×200点", "compute_kde_params.py", "technical-overview",
    "per param: {x_values[200], normal_density[200], fault_density[200], youdens_j}"],
   ["inventory_policy_optimized.csv", "CSV", "~20行(备件种类)", "decision", "inventory/index",
    "part_name, current_stock, safety_stock(s), target_stock(S), demand_rate, lead_time, unit_cost, stockout_risk"],
   ["strategy_comparison.csv", "CSV", "3行(3策略)", "decision", "index/home",
    "strategy, order_count, total_cost, avg_cost, total_downtime, quality_risk_index"],
   ["pareto_frontier.json", "JSON", "~200点", "decision(多目标优化)", "index",
    "per point: {cost, downtime, quality_risk, strategy_label}"],
   ["industrial_maintenance_plan.csv", "CSV", "~20-40行", "decision", "home/index",
    "machine_id, action, priority, technician, spare_parts, downtime_window, expected_savings, risk_before, risk_after"],
   ["spare_parts_plan.csv", "CSV", "~30行", "decision", "home/index",
    "part_name, part_number, quantity, unit_cost, supplier, lead_time_days, associated_machines"],
   ["technician_schedule.csv", "CSV", "~40行", "decision", "index/technicians",
    "technician_id, machine_id, task, start_time, end_time, duration_hours"],
   ["downtime_schedule.csv", "CSV", "~20行", "decision", "index",
    "machine_id, start_time, end_time, duration_hours, production_impact, conflict_flag"],
   ["procurement_orders.csv", "CSV", "~15行", "decision", "inventory",
    "order_id, part_name, quantity, status, order_date, estimated_arrival, supplier, total_cost"],
   ["log.csv", "CSV", "2999行", "原始数据(脱敏)", "index Tab1/Tab2/KDE计算",
    "Date, Equipment.Id, Failure.Equipment.Type, Op.Voltage, Op.Amperage, Op.Temperature, Rotor Speed"],
   ["backtest_lead_time_summary.csv", "CSV", "~30行", "stat_inference(回溯验证)", "index Tab5",
    "threshold_level, lead_time_bin, detected_count, miss_count"],
   ["backtest_by_fault_group.csv", "CSV", "~9行(9故障类型)", "stat_inference", "index Tab5",
    "fault_group, detection_rate, avg_lead_time, total_events"],
   ["degradation_status.json", "JSON", "1对象", "routes.py(实时更新)", "navbar/home",
    "current_level(FULL/STAT_ONLY/RULE_ONLY/EMERGENCY), reason, timestamp"],
   ["workflow_state.db", "SQLite", "—", "workflows(APScheduler)", "workflows",
    "job_id, last_run, last_status, last_duration, result_summary"]],
  col_widths=[4.0, 1.0, 2.5, 3.0, 3.5, 3.5])

PAGE_BREAK()

doc.add_heading('附录B：API端点完整列表', level=1)
T(["端点", "方法", "功能说明", "请求示例/参数"],
  [["/api/chat", "POST(SSE)", "AI Copilot流式对话——发送用户消息，接收SSE事件流(text_delta/tool_call/chart/report/rag_citations/done)", "Body: {message, history, role}"],
   ["/api/maintenance/strategy", "POST", "切换维护策略——后端重新执行决策引擎，生成新工单/备件/调度CSV", "Body: {strategy: 'quality_first'}"],
   ["/api/maintenance/machines-summary", "GET", "获取100台设备的预计算诊断摘要(~15KB)，避免前端加载900KB z_scores.csv", "—"],
   ["/api/health-summary", "GET", "全局健康摘要——关键设备数/最低分/Top5最低健康分设备等", "—"],
   ["/api/assistant/explain", "POST(SSE)", "评委讲解AI回退——话术库未命中时生成定制讲解", "Body: {page, section, context, question}"],
   ["/api/fault-injection", "POST", "故障注入演示——在内存中模拟传感器异常，不写文件", "Body: {machine_id, fault_type, severity}"],
   ["/api/work-order-tracking/list", "GET", "工单Kanban全量列表——所有状态的工单", "Query: ?strategy=all&status=pending"],
   ["/api/work-order-tracking/detail/<id>", "GET", "单个工单完整详情——含状态历史时间线", "Path: machine_id"],
   ["/api/work-order/create", "POST", "创建新工单并加入跟踪看板", "Body: {machine_id, priority, action}"],
   ["/api/work-order/update-status", "POST", "工单状态变更——6状态机流转", "Body: {machine_id, new_status, note}"],
   ["/api/technicians", "GET/POST", "获取技师列表 / 新增技师", "GET: — | POST: {name, type, email, phone, max_load, skills}"],
   ["/api/technicians/<id>", "PUT/DELETE", "更新/删除技师信息", "PUT: {name, type, ...} | DELETE: —"],
   ["/api/technicians/assign", "POST", "为工单分配技师（支持多选）", "Body: {machine_id, technician_ids: [...]}"],
   ["/api/technicians/workloads/summary", "GET", "技师负载汇总——总人数/可用/忙碌/超载", "—"],
   ["/api/inventory/stock", "GET", "库存状态——含(s,S)对比和缺口计算", "—"],
   ["/api/inventory/procurement", "GET", "采购订单列表——含状态跟踪", "—"],
   ["/api/inventory/procurement/generate", "POST", "基于缺口分析生成采购建议（EOQ计算）", "—"],
   ["/api/inventory/procurement/update-status", "POST", "更新采购订单状态（5状态流转）", "Body: {order_id, new_status}"],
   ["/api/inventory/adjust", "POST", "库存数量调整——记录变更日志", "Body: {part_name, delta, reason}"],
   ["/api/workflows/status", "GET", "三个定时任务的当前状态", "—"],
   ["/api/workflows/trigger/<jobId>", "POST", "手动触发定时任务", "Path: wo_timeout_check / daily_health_check / weekly_report"],
   ["/api/workflows/history", "GET", "定时任务执行历史——可筛选", "Query: ?job=weekly_report"],
   ["/api/knowledge-base/stats", "GET", "知识库统计——三个集合的文档数/分块数/嵌入状态", "—"],
   ["/api/knowledge-base/search", "GET", "知识库语义检索——基于ChromaDB+BGE", "Query: ?q=健康评分公式&collection=all"],
   ["/api/knowledge-base/degrade-test", "POST", "降级对比测试——三列并排展示BGE/DeepSeek/TF-IDF检索结果", "Body: {query}"],
   ["/api/reports", "GET", "报告列表——可按类型/关键词筛选排序", "Query: ?type=weekly&search=2026"],
   ["/api/reports/generate-pdf", "POST", "HTML报告→PDF转换（WeasyPrint）", "Body: {filename}"],
   ["/api/reports/delete", "POST", "删除报告文件", "Body: {filename}"]],
  col_widths=[4.5, 2.0, 5.5, 5.0])

PAGE_BREAK()

doc.add_heading('附录C：技术栈总表', level=1)
T(["层级", "技术", "版本", "用途"],
  [["后端框架", "Python + FastAPI", "3.10+ / 0.104+", "Web服务、REST API、SSE流式响应、中间件"],
   ["AI大模型", "DeepSeek (可配置)", "deepseek-chat", "AI对话、RAG语义理解、评委讲解、Tool Calling"],
   ["嵌入模型", "BGE (BAAI/bge-large-zh-v1.5)", "1024维", "中文文档语义向量化、ChromaDB底层编码器"],
   ["向量数据库", "ChromaDB", "0.4+", "RAG知识库存储、语义相似度检索、元数据过滤"],
   ["ML框架", "XGBoost + scikit-learn", "2.0+ / 1.3+", "异常检测模型训练、交叉验证、算法对比"],
   ["可解释AI", "SHAP (TreeExplainer)", "0.42+", "特征归因、全局/局部可解释性、瀑布图/条形图"],
   ["科学计算", "pandas + numpy + scipy", "—", "数据处理、统计检验(KDE/Z-Score/T²/ANOVA)、Bootstrap"],
   ["任务调度", "APScheduler", "3.10+", "定时任务（每日健康检查/超时巡检/周报生成）"],
   ["邮件", "smtplib + email.mime", "标准库", "告警升级通知 + 日报/周报投递"],
   ["PDF生成", "WeasyPrint", "—", "HTML报告→PDF（支持CSS打印样式）"],
   ["前端渲染", "Vanilla JavaScript (ES6+)", "—", "原生DOM操作、无框架依赖、模块化设计"],
   ["图表库", "ECharts + ECharts GL", "5.5.0 / 2.0.9", "2D交互图表(散点/柱状/折线/饼图/雷达/热力图/桑基图) + 3D WebGL图表"],
   ["轻量图表", "Chart.js", "4.4.0", "库存管理页面的柱状图和环形图"],
   ["3D引擎", "Three.js (ES Module)", "0.160.0", "鹰眼球体3D场景、Canvas纹理、粒子系统"],
   ["手势识别", "MediaPipe HandLandmarker", "0.10.14", "21点手部关键点检测 → 3D球体操控"],
   ["CSV解析", "PapaParse", "5.4.1", "前端CSV流式解析(最大900KB z_scores.csv)"],
   ["Markdown渲染", "marked.js", "latest", "AI对话消息的Markdown→HTML实时渲染"],
   ["CSS设计系统", "iOS26玻璃态 + CSS自定义属性", "—", "navbar.js注入的全局设计令牌、暗/亮双主题"],
   ["Python文档生成", "python-docx", "—", "本手册的.docx生成"]],
  col_widths=[2.5, 5.0, 2.0, 7.0])

PAGE_BREAK()

doc.add_heading('附录D：关键术语表', level=1)
T(["术语", "全称/解释", "出现场景"],
  [["CNC", "Computer Numerical Control（计算机数控机床）", "平台监控对象——100台编号CNC_001~100"],
   ["Z-Score", "标准化得分 Z = (x-μ)/σ，衡量当前值偏离历史均值的程度（以标准差为单位）", "单传感器异常检测的统计基础，所有图表中的\"Z值\""],
   ["复合Z-Score", "Z_comp = √(Σw²Z²)/√(Σw²)，多传感器Z-Score的加权融合", "设备健康评分的主要输入之一"],
   ["Hotelling T²", "多元马氏距离推广 T² = n(x̄-μ)'S⁻¹(x̄-μ)，考虑传感器协方差", "多传感器联合异常检测——捕获传感器间的相关异常"],
   ["SHAP", "SHapley Additive exPlanations，基于合作博弈Shapley值的特征归因方法", "每个告警的\"为什么\"——哪个传感器参数贡献了多少异常分数"],
   ["Youden's J", "J = Sensitivity + Specificity - 1 = TPR - FPR，综合二分类能力指标", "单参数检测能力上限评估——KDE天花板Demo的核心指标"],
   ["KDE", "Kernel Density Estimation（核密度估计），非参数概率密度估计方法", "传感器参数正常vs故障分布的可视化——论证\"信息天花板\""],
   ["RAG", "Retrieval-Augmented Generation（检索增强生成）", "AI Copilot的知识库问答——\"先查资料再回答\""],
   ["SSE", "Server-Sent Events（服务器推送事件）", "AI对话的流式文本输出——打字机效果"],
   ["DAG", "Directed Acyclic Graph（有向无环图）", "五层流水线的任务依赖与并行调度"],
   ["(s,S)策略", "再订购点s + 目标库存水平S 的两级库存策略", "备件库存管理的核心数学模型"],
   ["EOQ", "Economic Order Quantity（经济订购量）= √(2DS/H)", "单次最优采购批量——平衡订购成本和持有成本"],
   ["SLA", "Service Level Agreement（服务水平协议）", "工单完成的时限管理——超时自动升级"],
   ["P1/P2/P3", "Priority 1/2/3（三级优先级）", "工单紧急程度标注——P1需立即处理"],
   ["MCP", "Model Context Protocol（模型上下文协议）", "AI助手调用外部工具的标准化协议——27个注册工具"],
   ["Pareto前沿", "多目标优化中无法在不损害其他目标的前提下继续改进的解集合", "成本-停机时间-质量风险的三维最优边界"],
   ["CV%", "Coefficient of Variation（变异系数）= σ/μ × 100%", "传感器参数不稳定性度量——反映设备运转的稳定性"],
   ["RUL", "Remaining Useful Life（剩余可用寿命），单位天", "设备还能正常运行多少天——基于退化趋势的Weibull估计"],
   ["Wilson 95% CI", "Wilson Score置信区间——对小样本比例更准确的区间估计", "桑基图中的条件概率标注——避免小样本误导"]],
  col_widths=[2.5, 6.5, 6.5])

PAGE_BREAK()

doc.add_heading('附录E：颜色编码参考', level=1)
P("平台所有图表和组件使用统一的颜色语义编码。以下为跨页面的颜色编码标准。")
T(["语义色", "CSS变量名", "十六进制", "使用含义"],
  [["Primary Cyan", "accent-cyan", "#66d9c8", "主色调——激活状态、当前选中、图表系列1、正常/健康标识"],
   ["Amber/Warning", "accent-amber", "#ffb340", "警告状态——P2优先级、图表系列2、库存不足、过渡状态"],
   ["Red/Critical", "accent-red", "#ff453a", "危险状态——P1优先级、健康评分<40、告警ALARM、删除操作"],
   ["Green/Success", "accent-green", "#30d158", "成功/健康——健康评分≥85、已完成状态、节约效益、图表系列3"],
   ["Blue/Info", "accent-blue", "#6db5f9", "信息——P3优先级、图表系列4、链接、正常运转"],
   ["Purple/Developer", "accent-purple", "#bf5af2", "开发者——RAG工具卡片、运筹优化、开发者角色标识"],
   ["Grey/Neutral", "surface/glass", "#1c1c1e / rgba(255,255,255,0.05)", "背景、卡片、面板——iOS26玻璃态"],
   ["Health Gradient", "—", "绿→浅绿→琥珀→橙→红", "健康评分0-100的五级映射: ≥85绿, 70-84浅绿, 55-69琥珀, 40-54橙, <40红"]],
  col_widths=[2.5, 3.5, 2.5, 7.5])

# ── SAVE ──
doc.save(OUTPUT_PATH)
print(f"Document saved: {OUTPUT_PATH}")
print(f"File size: {os.path.getsize(OUTPUT_PATH):,} bytes")
